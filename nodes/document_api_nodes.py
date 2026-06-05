from __future__ import annotations

import json
import mimetypes
import os
import re
import time
import urllib.parse
from pathlib import Path
from typing import Any, Optional
from .generic_helpers import AlwaysDirtyMixin, SELECTABLE_RETURN_TYPES, SELECTABLE_RETURN_NAMES, \
    SELECTABLE_OUTPUT_IS_LIST, _session_request, _dirty_token, _json_text, _extract_data, _get_document_browser_root, \
    _resolve_document_file_path, ZMONGO_FILE_PATH, ZMONGO_FILENAME, ZMONGO_DOCUMENT_ID, _clean_scalar, \
    _coerce_file_path_link, _ensure_payload_dict, _as_comfy_list, _selectable_tail, _parse_json_object, _extract_doc_id, \
    _raise_if_failed, _extract_documents, _extract_doc_ids_from_documents, _extract_filenames_from_documents, \
    _coerce_document_id, _prefer_link_value, _extract_document, _extract_text, _dedupe_strings, _flatten_document_paths, \
    _indexed_list_text, _extract_field_paths, _coerce_field_path, _safe_get_by_path, _value_items, \
    _is_immutable_document_field_path, _blocked_immutable_field_payload, _parse_any_json, _parse_json_list, \
    _error_payload, ZMONGO_FIELD_PATH, _strip_index_prefix, _coerce_filename_link, _coerce_text_link, \
    _coerce_document_id_link, _document_summary, ZMONGO_TEXT, _read_file_as_base64, _list_document_files, \
    _session_get_doc, _session_api_request, _session_save_value

DEFAULT_DOCUMENT_PREFIX = os.getenv("ZMONGO_DOCUMENT_API_PREFIX", "/documents").strip().rstrip("/") or "/documents"
DEFAULT_DOCUMENT_COLLECTION = os.getenv("ZMONGO_DOCUMENT_COLLECTION", "documents").strip() or "documents"
DEFAULT_MAX_UPLOAD_BYTES = int(os.getenv("ZMONGO_DOCUMENT_NODE_MAX_UPLOAD_BYTES", str(64 * 1024 * 1024)))

DEFAULT_DOCUMENT_DISCOVERY_EXTENSIONS = os.getenv(
    "ZMONGO_DOCUMENT_BROWSER_EXTENSIONS",
    ".pdf,.docx,.txt,.md,.csv,.json,.log,.rtf",
).strip() or ".pdf,.docx,.txt,.md,.csv,.json,.log,.rtf"
DEFAULT_DOCUMENT_DISCOVERY_MAX_RESULTS = int(os.getenv("ZMONGO_DOCUMENT_BROWSER_MAX_RESULTS", "500"))
DEFAULT_DOCUMENT_DISCOVERY_MAX_DEPTH = int(os.getenv("ZMONGO_DOCUMENT_BROWSER_MAX_DEPTH", "6"))

DOCUMENT_BROWSER_ROOT_MODES = [
    "Documents",
    "Comfy Input",
    "Home",
    "Current Working Directory",
    "Filesystem Root",
    "Custom Root",
]


def _split_extensions(extensions: str) -> set[str]:
    cleaned: set[str] = set()
    for item in re.split(r"[,;\s]+", str(extensions or "")):
        item = item.strip().lower()
        if not item:
            continue
        if item in {"*", ".*", "all"}:
            return set()
        if not item.startswith("."):
            item = "." + item
        cleaned.add(item)
    return cleaned


def _safe_resolve_path(value: Any) -> Path:
    raw = _clean_scalar(value)
    if not raw:
        return Path.home().resolve()
    return Path(os.path.expandvars(os.path.expanduser(raw))).resolve()


def _document_browser_root_for_mode(root_mode: str, custom_root: str = "") -> Path:
    mode = str(root_mode or "Documents").strip()

    if mode == "Home":
        return Path.home().resolve()

    if mode == "Current Working Directory":
        return Path.cwd().resolve()

    if mode == "Filesystem Root":
        return Path(Path.home().anchor or "/").resolve()

    if mode == "Custom Root":
        return _safe_resolve_path(custom_root)

    try:
        return Path(_get_document_browser_root(root_mode=mode, custom_root=custom_root)).resolve()
    except Exception:
        if mode == "Comfy Input" and folder_paths is not None:
            try:
                return Path(folder_paths.get_input_directory()).resolve()
            except Exception:
                pass
        return (Path.home() / "Documents").resolve()


def _is_probably_hidden(path: Path) -> bool:
    return any(part.startswith(".") for part in path.parts if part not in {path.anchor, "/"})


def _should_skip_discovery_dir(path: Path, include_hidden: bool) -> bool:
    name = path.name
    if not include_hidden and name.startswith("."):
        return True
    return name in {
        "__pycache__",
        ".git",
        ".hg",
        ".svn",
        "node_modules",
        ".venv",
        "venv",
        "env",
        "site-packages",
        "dist-packages",
        "proc",
        "sys",
        "dev",
        "run",
        "tmp",
        "var",
    }


def _discover_document_files(
    *,
    root: Path,
    extensions: str = DEFAULT_DOCUMENT_DISCOVERY_EXTENSIONS,
    recursive: bool = True,
    max_results: int = DEFAULT_DOCUMENT_DISCOVERY_MAX_RESULTS,
    max_depth: int = DEFAULT_DOCUMENT_DISCOVERY_MAX_DEPTH,
    search_text: str = "",
    include_hidden: bool = False,
) -> list[Path]:
    root = Path(root).resolve()
    allowed_exts = _split_extensions(extensions)
    limit = max(1, min(int(max_results or DEFAULT_DOCUMENT_DISCOVERY_MAX_RESULTS), 10000))
    depth_limit = max(0, min(int(max_depth or DEFAULT_DOCUMENT_DISCOVERY_MAX_DEPTH), 25))
    needle = str(search_text or "").strip().lower()

    if not root.exists():
        return []

    if root.is_file():
        if allowed_exts and root.suffix.lower() not in allowed_exts:
            return []
        if needle and needle not in root.name.lower() and needle not in str(root).lower():
            return []
        return [root]

    found: list[Path] = []
    stack: list[tuple[Path, int]] = [(root, 0)]
    seen_dirs: set[str] = set()

    while stack and len(found) < limit:
        current, depth = stack.pop()
        try:
            current_key = str(current.resolve())
        except Exception:
            current_key = str(current)
        if current_key in seen_dirs:
            continue
        seen_dirs.add(current_key)

        try:
            entries = sorted(current.iterdir(), key=lambda p: (not p.is_dir(), str(p).lower()))
        except (PermissionError, OSError):
            continue

        for entry in entries:
            if len(found) >= limit:
                break

            try:
                if entry.is_dir():
                    if recursive and depth < depth_limit and not _should_skip_discovery_dir(entry, include_hidden):
                        stack.append((entry, depth + 1))
                    continue

                if not entry.is_file():
                    continue

                if not include_hidden and _is_probably_hidden(entry):
                    continue

                if allowed_exts and entry.suffix.lower() not in allowed_exts:
                    continue

                entry_text = f"{entry.name}\n{entry}".lower()
                if needle and needle not in entry_text:
                    continue

                found.append(entry.resolve())
            except (PermissionError, OSError):
                continue

    found.sort(key=lambda p: str(p).lower())
    return found


def _discovered_files_for_dropdown() -> list[str]:
    roots: list[Path] = []
    for mode in ("Documents", "Comfy Input", "Home"):
        try:
            root = _document_browser_root_for_mode(mode)
            if root.exists() and root not in roots:
                roots.append(root)
        except Exception:
            continue

    values: list[str] = [""]
    seen: set[str] = set()
    per_root_limit = max(25, min(DEFAULT_DOCUMENT_DISCOVERY_MAX_RESULTS, 250))
    for root in roots:
        for file_path in _discover_document_files(
            root=root,
            recursive=True,
            max_results=per_root_limit,
            max_depth=DEFAULT_DOCUMENT_DISCOVERY_MAX_DEPTH,
        ):
            value = str(file_path)
            if value in seen:
                continue
            seen.add(value)
            values.append(value)

    return values or [""]



def _is_route_not_found(payload: Any) -> bool:
    payload = _ensure_payload_dict(payload)
    status_code = int(payload.get("status_code") or 0)
    if status_code == 404:
        return True
    raw_text = str(payload.get("raw_text") or "")
    message = str(payload.get("message") or "")
    return "404 Not Found" in raw_text or "Not Found" == message.strip()


def _generic_document_get_fallback(session: Any, document_id: str, *, collection_name: str = DEFAULT_DOCUMENT_COLLECTION) -> dict[str, Any]:
    if not document_id:
        return _ensure_payload_dict({
            "success": False,
            "message": "Document get fallback skipped: document_id is empty.",
            "data": {"collection": collection_name, "document_id": document_id},
            "error": {"msg": "missing_document_id"},
            "status_code": 400,
        })
    return _session_get_doc(session, collection_name, document_id, cache=False)


def _generic_document_query_fallback(
    session: Any,
    *,
    query: dict[str, Any],
    sort: list[Any],
    limit: int,
    skip: int,
    collection_name: str = DEFAULT_DOCUMENT_COLLECTION,
) -> dict[str, Any]:
    list_docs = getattr(session, "list_docs", None)
    if callable(list_docs):
        try:
            return _ensure_payload_dict(
                list_docs(collection=collection_name, query=query, limit=limit, skip=skip)
            )
        except TypeError:
            pass
        except Exception as exc:
            return _ensure_payload_dict({
                "success": False,
                "message": f"Generic document list fallback failed: {exc}",
                "data": {"collection": collection_name, "query": query, "limit": limit, "skip": skip},
                "error": {"type": exc.__class__.__name__, "msg": str(exc)},
                "status_code": 0,
            })

    return _session_api_request(
        session,
        "POST",
        "/api/query",
        json_body={
            "collection": collection_name,
            "query": query,
            "sort": sort,
            "limit": int(limit),
            "skip": int(skip),
            "many": True,
            "cache": False,
        },
    )


def _generic_document_create_fallback(
    session: Any,
    *,
    document: dict[str, Any],
    collection_name: str = DEFAULT_DOCUMENT_COLLECTION,
) -> dict[str, Any]:
    create_doc = getattr(session, "create_doc", None)
    if callable(create_doc):
        try:
            return _ensure_payload_dict(create_doc(collection=collection_name, document=document))
        except Exception as exc:
            return _ensure_payload_dict({
                "success": False,
                "message": f"Generic document create fallback failed: {exc}",
                "data": {"collection": collection_name, "document": document},
                "error": {"type": exc.__class__.__name__, "msg": str(exc)},
                "status_code": 0,
            })

    return _session_api_request(
        session,
        "POST",
        "/api/doc/create",
        json_body={"collection": collection_name, "document": document},
    )


# -----------------------------------------------------------------------------
# Local text-document ingestion helpers
# -----------------------------------------------------------------------------

TEXT_DOCUMENT_SCHEMA_KIND = "zmongo_text_document"
TEXT_DOCUMENT_SCHEMA_VERSION = "1.0.0"
DEFAULT_TEXT_DOCUMENT_FIELD_ROOT = "document_text"
DEFAULT_TEXT_DOCUMENT_COLLECTION = os.getenv("ZMONGO_TEXT_DOCUMENT_COLLECTION", "text_documents").strip() or "text_documents"
DEFAULT_LOCAL_TEXT_DOCUMENT_MAX_BYTES = int(os.getenv("ZMONGO_LOCAL_TEXT_DOCUMENT_MAX_BYTES", str(256 * 1024)))


def _is_local_storage_session(session: Any) -> bool:
    return str(getattr(session, "storage_backend", "") or "").strip() == "local_file_store"


def _utf8_size(value: Any) -> int:
    return len(str(value or "").encode("utf-8"))


def _local_storage_limit_payload(*, path: Path, filename: str, text_size_bytes: int, refresh: str) -> dict[str, Any]:
    return _ensure_payload_dict({
        "success": False,
        "message": (
            f"Local File Store is a small CRUD store and will not save extracted text larger than "
            f"{DEFAULT_LOCAL_TEXT_DOCUMENT_MAX_BYTES} bytes. Use BusinessProcessApplications.com / hosted ZMongo "
            "for large files, where ZEmbedder.py handles document chunking and embeddings."
        ),
        "data": {
            "effective_file_path": str(path),
            "filename": filename,
            "text_size_bytes": int(text_size_bytes),
            "local_max_text_document_bytes": int(DEFAULT_LOCAL_TEXT_DOCUMENT_MAX_BYTES),
            "storage_backend": "local_file_store",
            "recommended_backend": "https://businessprocessapplications.com",
            "recommended_pipeline": "Upload the original file to hosted ZMongo; let backend ZEmbedder.py chunk/embed it.",
            "refresh": refresh,
        },
        "error": {"type": "LocalStorageItemTooLarge", "msg": "local_text_document_limit_exceeded"},
        "status_code": 413,
    })


def _clean_dot_path_segment(value: Any, fallback: str = "field") -> str:
    cleaned = re.sub(r"[^A-Za-z0-9_]+", "_", str(value or "").strip()).strip("_")
    return cleaned or fallback


def _safe_read_text_file(path: Path) -> str:
    raw = path.read_bytes()
    for encoding in ("utf-8", "utf-8-sig", "cp1252", "latin-1"):
        try:
            return raw.decode(encoding)
        except Exception:
            continue
    return raw.decode("utf-8", errors="replace")


def _extract_text_from_pdf_file(path: Path) -> str:
    try:
        import fitz  # PyMuPDF
    except Exception as exc:
        raise RuntimeError(
            "PDF text extraction requires PyMuPDF. Install it in the ComfyUI Python environment with: pip install pymupdf"
        ) from exc

    pieces: list[str] = []
    with fitz.open(str(path)) as doc:
        for page_index, page in enumerate(doc):
            page_text = page.get_text("text") or ""
            if page_text.strip():
                pieces.append(f"\n\n--- Page {page_index + 1} ---\n{page_text.strip()}")
    return "".join(pieces).strip()


def _extract_text_from_docx_file(path: Path) -> str:
    try:
        from docx import Document
    except Exception as exc:
        raise RuntimeError(
            "DOCX text extraction requires python-docx. Install it in the ComfyUI Python environment with: pip install python-docx"
        ) from exc

    document = Document(str(path))
    pieces: list[str] = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            pieces.append(text)

    for table_index, table in enumerate(document.tables):
        rows: list[str] = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            row_text = "\t".join(cells).strip()
            if row_text:
                rows.append(row_text)
        if rows:
            pieces.append(f"\n[Table {table_index + 1}]\n" + "\n".join(rows))

    return "\n\n".join(pieces).strip()


def _extract_text_from_supported_document(path: Path) -> tuple[str, str]:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return _extract_text_from_pdf_file(path), "pymupdf"
    if suffix == ".docx":
        return _extract_text_from_docx_file(path), "python-docx"
    if suffix in {".txt", ".md", ".markdown", ".csv", ".json", ".log"}:
        return _safe_read_text_file(path), "plain-text"
    raise ValueError(f"Unsupported text-ingestion file type: {suffix}. Supported: .pdf, .docx, .txt, .md, .csv, .json, .log")



def _text_quality_score(value: Any) -> int:
    """
    Score likely extracted document text. Higher is better.

    This intentionally penalizes short strings containing control/binary-looking
    characters because local document APIs can sometimes expose file bytes or a
    stale/corrupt top-level ``text`` value.  Long readable text from
    ``document_text.raw_text`` should win over a short binary-looking ``text``.
    """
    if not isinstance(value, str):
        return -1

    text = value.strip()
    if not text:
        return 0

    length = len(text)
    printable = sum(1 for ch in text if ch.isprintable() or ch in "\r\n\t")
    controls = sum(1 for ch in text if (ord(ch) < 32 and ch not in "\r\n\t"))
    replacement = text.count("\ufffd")
    whitespace = sum(1 for ch in text if ch.isspace())
    letters = sum(1 for ch in text if ch.isalpha())

    printable_ratio = printable / max(1, length)
    letter_ratio = letters / max(1, length)
    whitespace_ratio = whitespace / max(1, length)

    score = min(length, 100000)
    score += int(printable_ratio * 1000)
    score += int(letter_ratio * 500)
    score += int(whitespace_ratio * 250)
    score -= controls * 1000
    score -= replacement * 500

    if length < 32 and (controls or printable_ratio < 0.85 or letter_ratio < 0.25):
        score -= 10000

    return score


def _looks_like_bad_text_value(value: Any) -> bool:
    if not isinstance(value, str):
        return False
    text = value.strip()
    if not text:
        return False
    length = len(text)
    printable = sum(1 for ch in text if ch.isprintable() or ch in "\r\n\t")
    controls = sum(1 for ch in text if (ord(ch) < 32 and ch not in "\r\n\t"))
    letters = sum(1 for ch in text if ch.isalpha())
    return (
        (length < 64 and controls > 0)
        or (length > 0 and printable / max(1, length) < 0.85)
        or (length < 64 and letters / max(1, length) < 0.20 and controls > 0)
    )


def _best_text_from_document(document: Any) -> tuple[str, str, int]:
    """Return the best text-like field from a document plus its source path."""
    if not isinstance(document, dict):
        return "", "", 0

    candidates: list[tuple[str, Any]] = []
    for candidate_path in (
        "document_text.raw_text",
        "document_text.text",
        "content.text",
        "extracted_text",
        "raw_text",
        "text",
    ):
        marker = object()
        value = _safe_get_by_path(document, candidate_path, marker)
        if value is not marker:
            candidates.append((candidate_path, value))

    best_text = ""
    best_path = ""
    best_score = -1
    for source_path, value in candidates:
        if not isinstance(value, str):
            continue
        score = _text_quality_score(value)
        if score > best_score:
            best_text = value
            best_path = source_path
            best_score = score

    return best_text, best_path, best_score


def _query_existing_text_document(
    session: Any,
    *,
    collection_name: str,
    document_name: str,
) -> tuple[dict[str, Any], str, dict[str, Any]]:
    query = {"schema_kind": TEXT_DOCUMENT_SCHEMA_KIND, "document_name": document_name}
    payload = _generic_document_query_fallback(
        session,
        query=query,
        sort=[["updated_at_unix", -1]],
        limit=1,
        skip=0,
        collection_name=collection_name,
    )
    documents = _extract_documents(payload)
    document = documents[0] if documents else {}
    document_id = _coerce_document_id(document) if document else ""
    return document, document_id, payload


def _save_text_document_by_dot_paths(
    session: Any,
    *,
    collection_name: str,
    document_id: str,
    document: dict[str, Any],
    field_root: str,
) -> dict[str, Any]:
    failures: list[str] = []
    saved_paths: list[str] = []

    root = _clean_dot_path_segment(field_root, DEFAULT_TEXT_DOCUMENT_FIELD_ROOT)
    save_items = {
        "schema_kind": document.get("schema_kind"),
        "schema_version": document.get("schema_version"),
        "document_name": document.get("document_name"),
        "collection_name": document.get("collection_name"),
        "source_file": document.get("source_file") or {},
        "metadata": document.get("metadata") or {},
        root: document.get(root) or {},
        "updated_at_unix": document.get("updated_at_unix"),
    }

    for field_path, value in save_items.items():
        payload = _session_save_value(
            session,
            collection_name=collection_name,
            document_id=document_id,
            query=None,
            field_path=field_path,
            value=value,
            upsert=False,
        )
        if payload.get("success"):
            saved_paths.append(field_path)
        else:
            failures.append(f"{field_path}: {payload.get('message') or payload.get('error')}")

    return _ensure_payload_dict({
        "success": not failures,
        "message": "Updated existing ingested text document." if not failures else "Some text document field writes failed.",
        "data": {"document_id": document_id, "saved_paths": saved_paths, "failures": failures},
        "error": {"msg": "; ".join(failures[:5])} if failures else None,
        "status_code": 200 if not failures else 500,
    })

try:
    import folder_paths  # ComfyUI runtime helper for input/output/temp directories
except Exception:  # Allows py_compile outside ComfyUI
    folder_paths = None

# -----------------------------------------------------------------------------
# Nodes
# -----------------------------------------------------------------------------

class ZMongoDocumentHealth(AlwaysDirtyMixin):
    @classmethod
    def INPUT_TYPES(cls):
        return {"required": {"session": ("ZMONGO_API_SESSION",), "document_prefix": ("STRING", {"default": DEFAULT_DOCUMENT_PREFIX})}, "optional": {"refresh_token": ("STRING", {"default": ""})}}

    RETURN_TYPES = ("STRING", "BOOLEAN", "STRING") + SELECTABLE_RETURN_TYPES
    RETURN_NAMES = ("json", "success", "refresh") + SELECTABLE_RETURN_NAMES
    OUTPUT_IS_LIST = (False, False, False) + SELECTABLE_OUTPUT_IS_LIST
    FUNCTION = "health"
    CATEGORY = "ZMongo/03 Documents"
    OUTPUT_NODE = True

    def health(self, session: Any, document_prefix: str, refresh_token: str = ""):
        payload = _session_request(session, "GET", document_prefix, "/health")
        refresh = _dirty_token("document_health", refresh_token)
        return (_json_text(payload), bool(payload.get("success")), refresh, *_selectable_tail([payload.get("message") or "health"]))


class ZMongoDocumentWhoAmI(AlwaysDirtyMixin):
    @classmethod
    def INPUT_TYPES(cls):
        return {"required": {"session": ("ZMONGO_API_SESSION",), "document_prefix": ("STRING", {"default": DEFAULT_DOCUMENT_PREFIX})}, "optional": {"refresh_token": ("STRING", {"default": ""})}}

    RETURN_TYPES = ("STRING", "STRING", "STRING", "BOOLEAN", "STRING") + SELECTABLE_RETURN_TYPES
    RETURN_NAMES = ("json", "username", "silo_db_name", "success", "refresh") + SELECTABLE_RETURN_NAMES
    OUTPUT_IS_LIST = (False, False, False, False, False) + SELECTABLE_OUTPUT_IS_LIST
    FUNCTION = "whoami"
    CATEGORY = "ZMongo/03 Documents"
    OUTPUT_NODE = True

    def whoami(self, session: Any, document_prefix: str, refresh_token: str = ""):
        payload = _session_request(session, "GET", document_prefix, "/api/whoami")
        data = _extract_data(payload)
        username = str(data.get("username") or "")
        silo_db_name = str(data.get("silo_db_name") or data.get("db_name") or "")
        return (_json_text(payload), username, silo_db_name, bool(payload.get("success")), _dirty_token("document_whoami", refresh_token), *_selectable_tail([username]))


class ZMongoDocumentFilePathBrowser(AlwaysDirtyMixin):
    @classmethod
    def INPUT_TYPES(cls):
        return {
            "required": {
                "root_mode": (DOCUMENT_BROWSER_ROOT_MODES, {"default": "Documents"}),
                "selected_file": (_discovered_files_for_dropdown(),),
                "manual_path": ("STRING", {"default": ""}),
                "custom_root": ("STRING", {"default": ""}),
                "recursive": ("BOOLEAN", {"default": True}),
                "extensions": (
                    "STRING",
                    {
                        "default": DEFAULT_DOCUMENT_DISCOVERY_EXTENSIONS,
                        "tooltip": "Comma/space-separated extensions. Use * for all files.",
                    },
                ),
                "search_text": (
                    "STRING",
                    {
                        "default": "",
                        "tooltip": "Optional case-insensitive filename/path filter.",
                    },
                ),
                "max_results": (
                    "INT",
                    {"default": DEFAULT_DOCUMENT_DISCOVERY_MAX_RESULTS, "min": 1, "max": 10000},
                ),
                "max_depth": (
                    "INT",
                    {"default": DEFAULT_DOCUMENT_DISCOVERY_MAX_DEPTH, "min": 0, "max": 25},
                ),
                "include_hidden": ("BOOLEAN", {"default": False}),
            },
            "optional": {"refresh_token": ("STRING", {"default": ""})},
        }

    RETURN_TYPES = (
        "STRING",
        ZMONGO_FILE_PATH,
        "STRING",
        ZMONGO_FILENAME,
        "STRING",
        "INT",
        "BOOLEAN",
        "STRING",
        "STRING",
        "*",
        "*",
        "STRING",
        "INT",
        "STRING",
    ) + SELECTABLE_RETURN_TYPES
    RETURN_NAMES = (
        "file_path",
        "file_path_link",
        "filename",
        "filename_link",
        "content_type",
        "size_bytes",
        "exists",
        "browser_root",
        "refresh",
        "file_paths",
        "filenames",
        "indexed_file_paths",
        "discovered_count",
        "status",
    ) + SELECTABLE_RETURN_NAMES
    OUTPUT_IS_LIST = (False, False, False, False, False, False, False, False, False, True, True, False, False, False) + SELECTABLE_OUTPUT_IS_LIST
    FUNCTION = "resolve_file_path"
    CATEGORY = "ZMongo/03 Documents"
    OUTPUT_NODE = True

    @staticmethod
    def _resolve_selected_file(selected_file: str, manual_path: str, browser_root: Path) -> Path:
        manual = _clean_scalar(manual_path)
        selected = _clean_scalar(selected_file)

        if manual:
            return _safe_resolve_path(manual)

        if selected:
            selected_path = Path(os.path.expandvars(os.path.expanduser(selected)))
            if selected_path.is_absolute():
                return selected_path.resolve()
            return (browser_root / selected_path).resolve()

        return browser_root

    def resolve_file_path(
        self,
        root_mode: str = "Documents",
        selected_file: str = "",
        manual_path: str = "",
        custom_root: str = "",
        recursive: bool = True,
        extensions: str = DEFAULT_DOCUMENT_DISCOVERY_EXTENSIONS,
        search_text: str = "",
        max_results: int = DEFAULT_DOCUMENT_DISCOVERY_MAX_RESULTS,
        max_depth: int = DEFAULT_DOCUMENT_DISCOVERY_MAX_DEPTH,
        include_hidden: bool = False,
        refresh_token: str = "",
    ):
        refresh = _dirty_token(
            "document_file_browser",
            refresh_token,
            root_mode,
            selected_file,
            manual_path,
            custom_root,
            recursive,
            extensions,
            search_text,
            max_results,
            max_depth,
            include_hidden,
        )

        try:
            browser_root = _document_browser_root_for_mode(root_mode=root_mode, custom_root=custom_root)
            discovered = _discover_document_files(
                root=browser_root,
                extensions=extensions,
                recursive=bool(recursive),
                max_results=int(max_results),
                max_depth=int(max_depth),
                search_text=search_text,
                include_hidden=bool(include_hidden),
            )

            resolved = self._resolve_selected_file(selected_file, manual_path, browser_root)
            exists = resolved.exists() and resolved.is_file()

            # If the user did not select/type a file, expose the first discovered file
            # as the primary scalar while still returning the complete discovered list.
            if not exists and not _clean_scalar(manual_path) and not _clean_scalar(selected_file) and discovered:
                resolved = discovered[0]
                exists = True

            size_bytes = resolved.stat().st_size if exists else 0
            content_type = mimetypes.guess_type(str(resolved))[0] or "application/octet-stream"
            filename = resolved.name if exists else ""
            file_path = str(resolved) if exists else ""

            file_paths = [str(path) for path in discovered]
            if file_path and file_path not in file_paths:
                file_paths.insert(0, file_path)

            filenames = [Path(path).name for path in file_paths]
            indexed = _indexed_list_text(file_paths)

            payload = _ensure_payload_dict(
                {
                    "success": bool(exists or file_paths),
                    "message": f"Discovered {len(file_paths)} file(s) under {browser_root}.",
                    "data": {
                        "root_mode": root_mode,
                        "browser_root": str(browser_root),
                        "selected_file": _clean_scalar(selected_file),
                        "manual_path": _clean_scalar(manual_path),
                        "custom_root": _clean_scalar(custom_root),
                        "effective_file_path": file_path,
                        "exists": bool(exists),
                        "recursive": bool(recursive),
                        "extensions": sorted(_split_extensions(extensions)) or ["*"],
                        "search_text": _clean_scalar(search_text),
                        "max_results": int(max_results),
                        "max_depth": int(max_depth),
                        "include_hidden": bool(include_hidden),
                        "discovered_count": len(file_paths),
                        "file_paths": file_paths,
                        "filenames": filenames,
                        "refresh": refresh,
                    },
                    "error": None if (exists or file_paths) else {"msg": "No matching files discovered."},
                    "status_code": 200 if (exists or file_paths) else 404,
                }
            )

            status = payload.get("message") or ""
            return (
                file_path,
                file_path,
                filename,
                filename,
                content_type,
                int(size_bytes),
                bool(exists),
                str(browser_root),
                refresh,
                _as_comfy_list(file_paths),
                _as_comfy_list(filenames),
                indexed,
                len(file_paths),
                status,
                *_selectable_tail(file_paths),
            )
        except Exception as exc:
            try:
                browser_root = str(_document_browser_root_for_mode(root_mode=root_mode, custom_root=custom_root))
            except Exception:
                browser_root = ""
            payload = _ensure_payload_dict(
                {
                    "success": False,
                    "message": f"Document file discovery failed: {exc}",
                    "data": {
                        "root_mode": root_mode,
                        "browser_root": browser_root,
                        "selected_file": _clean_scalar(selected_file),
                        "manual_path": _clean_scalar(manual_path),
                        "custom_root": _clean_scalar(custom_root),
                        "refresh": refresh,
                    },
                    "error": {"type": exc.__class__.__name__, "msg": str(exc)},
                    "status_code": 0,
                }
            )
            return (
                "",
                "",
                "",
                "",
                "application/octet-stream",
                0,
                False,
                browser_root,
                refresh,
                _as_comfy_list([]),
                _as_comfy_list([]),
                _indexed_list_text([]),
                0,
                _json_text(payload),
                *_selectable_tail([]),
            )


class ZMongoDocumentUploadFile(AlwaysDirtyMixin):
    @classmethod
    def INPUT_TYPES(cls):
        return {"required": {"session": ("ZMONGO_API_SESSION",), "file_path": ("STRING", {"default": ""}), "case_id": ("STRING", {"default": ""}), "metadata_json": ("STRING", {"default": "{}", "multiline": True}), "document_prefix": ("STRING", {"default": DEFAULT_DOCUMENT_PREFIX})}, "optional": {"file_path_link": ("*",), "known_text": ("STRING", {"default": "", "multiline": True}), "refresh_token": ("STRING", {"default": ""})}}

    RETURN_TYPES = ("STRING", "STRING", ZMONGO_DOCUMENT_ID, "STRING", ZMONGO_FILENAME, "INT", "BOOLEAN", "STRING", "*", "*") + SELECTABLE_RETURN_TYPES
    RETURN_NAMES = ("json", "document_id", "document_id_link", "filename", "filename_link", "size_bytes", "success", "refresh", "document_ids", "filenames") + SELECTABLE_RETURN_NAMES
    OUTPUT_IS_LIST = (False, False, False, False, False, False, False, False, True, True) + SELECTABLE_OUTPUT_IS_LIST
    FUNCTION = "upload_file"
    CATEGORY = "ZMongo/03 Documents"
    OUTPUT_NODE = True

    def upload_file(self, session: Any, file_path: str, case_id: str, metadata_json: str, document_prefix: str, file_path_link: str = "", known_text: str = "", refresh_token: str = ""):
        refresh = _dirty_token("document_upload", refresh_token)

        primary_path = _clean_scalar(file_path)
        linked_path = _coerce_file_path_link(file_path_link)
        effective_file_path = linked_path or primary_path
        effective_path = Path(os.path.expanduser(effective_file_path or "")).resolve() if effective_file_path else None

        def nonfatal(payload: dict[str, Any], filename_hint: str = ""):
            name = filename_hint or (effective_path.name if effective_path is not None else "")
            return (
                _json_text(_ensure_payload_dict(payload)),
                "",
                "",
                name,
                name,
                0,
                False,
                refresh,
                _as_comfy_list([]),
                _as_comfy_list([name] if name else []),
                *_selectable_tail([]),
            )

        if effective_path is None or not effective_file_path:
            return nonfatal({
                "success": False,
                "message": "Upload skipped. No document file path was provided.",
                "data": {
                    "file_path": primary_path,
                    "file_path_link": _clean_scalar(file_path_link),
                    "effective_file_path": "",
                    "checks": [
                        "Connect 06 Document File Browser.file_path_link to 06 Upload Document File.file_path_link.",
                        "Select an actual file in the file browser node.",
                        "Do not connect browser_root, selected_file, indexed_items, or a directory path to file_path.",
                    ],
                    "refresh": refresh,
                },
                "error": {"type": "MissingFilePath", "msg": "No file path supplied."},
                "status_code": 400,
            })

        if effective_path.exists() and effective_path.is_dir():
            return nonfatal({
                "success": False,
                "message": f"Upload skipped. The selected path is a directory, not a document file: {effective_path}",
                "data": {
                    "file_path": primary_path,
                    "file_path_link": _clean_scalar(file_path_link),
                    "effective_file_path": str(effective_path),
                    "is_directory": True,
                    "checks": [
                        "Connect file_path_link from 06 Document File Browser, not browser_root.",
                        "In 06 Document File Browser, select a real PDF/TXT/DOCX/etc. file.",
                        "If using manual_path, enter the full path to a file, not a folder.",
                    ],
                    "refresh": refresh,
                },
                "error": {"type": "DirectorySelected", "msg": "Selected path is a directory."},
                "status_code": 400,
            }, effective_path.name)

        if not effective_path.exists() or not effective_path.is_file():
            return nonfatal({
                "success": False,
                "message": f"Upload skipped. Document file was not found: {effective_path}",
                "data": {
                    "file_path": primary_path,
                    "file_path_link": _clean_scalar(file_path_link),
                    "effective_file_path": str(effective_path),
                    "exists": effective_path.exists(),
                    "checks": [
                        "Confirm the file exists on the machine running ComfyUI, not merely on the browser client.",
                        "Put the document in /home/comfyuser/Documents or set ZMONGO_DOCUMENT_BROWSER_ROOT.",
                        "Reload the workflow if the file-browser dropdown has stale values.",
                    ],
                    "refresh": refresh,
                },
                "error": {"type": "FileNotFound", "msg": "Document file not found."},
                "status_code": 404,
            }, effective_path.name)

        try:
            filename, content_type, b64, size_bytes = _read_file_as_base64(str(effective_path))
            body = {
                "filename": filename,
                "content_type": content_type,
                "file_data": b64,
                "case_id": _clean_scalar(case_id) or None,
                "metadata": _parse_json_object(metadata_json, "metadata_json"),
                "text": known_text or "",
            }
            payload = _session_request(session, "POST", document_prefix, "/api/upload", json_body=body)
            data = _extract_data(payload)
            out_filename = str(data.get("filename") or filename)

            if not payload.get("success"):
                return (
                    _json_text(payload),
                    "",
                    "",
                    out_filename,
                    out_filename,
                    int(data.get("size_bytes") or size_bytes or 0),
                    False,
                    refresh,
                    _as_comfy_list([]),
                    _as_comfy_list([out_filename] if out_filename else []),
                    *_selectable_tail([]),
                )

            document_id = _extract_doc_id(payload)
            ids = [document_id] if document_id else []
            filenames = [out_filename] if out_filename else []
            return (
                _json_text(payload),
                document_id,
                document_id,
                out_filename,
                out_filename,
                int(data.get("size_bytes") or size_bytes or 0),
                bool(payload.get("success")),
                refresh,
                _as_comfy_list(ids),
                _as_comfy_list(filenames),
                *_selectable_tail(ids),
            )
        except Exception as exc:
            return nonfatal({
                "success": False,
                "message": f"Upload failed: {exc}",
                "data": {
                    "file_path": primary_path,
                    "file_path_link": _clean_scalar(file_path_link),
                    "effective_file_path": str(effective_path),
                    "refresh": refresh,
                },
                "error": {"type": exc.__class__.__name__, "msg": str(exc)},
                "status_code": 0,
            }, effective_path.name if effective_path is not None else "")



class ZMongoDocumentIngestTextFile(AlwaysDirtyMixin):
    """
    Pick a local text-based PDF, DOCX, or plain text file, extract text locally,
    and save it as one bounded ZMongo text-document record.

    Local File Store mode deliberately does not chunk documents. Large files
    should be uploaded to the hosted Business Process Applications / ZMongo
    backend, where ZEmbedder.py performs chunking and embedding.
    """

    @classmethod
    def INPUT_TYPES(cls):
        return {
            "required": {
                "session": ("ZMONGO_API_SESSION",),
                "file_path": ("STRING", {"default": ""}),
                "document_name": ("STRING", {"default": ""}),
                "collection_name": ("STRING", {"default": DEFAULT_TEXT_DOCUMENT_COLLECTION}),
                "field_root": ("STRING", {"default": DEFAULT_TEXT_DOCUMENT_FIELD_ROOT}),
                "overwrite": ("BOOLEAN", {"default": True}),
                "metadata_json": ("STRING", {"default": "{}", "multiline": True}),
                "tags_json": ("STRING", {"default": "[]", "multiline": True}),
            },
            "optional": {
                "file_path_link": ("*",),
                "refresh_token": ("STRING", {"default": ""}),
            },
        }

    RETURN_TYPES = (
        "STRING",
        "STRING",
        ZMONGO_DOCUMENT_ID,
        "STRING",
        ZMONGO_FILENAME,
        "STRING",
        "INT",
        "INT",
        "BOOLEAN",
        "STRING",
        "*",
        "*",
    ) + SELECTABLE_RETURN_TYPES
    RETURN_NAMES = (
        "json",
        "document_id",
        "document_id_link",
        "document_name",
        "filename_link",
        "text",
        "text_length",
        "local_limit_bytes",
        "success",
        "refresh",
        "document_ids",
        "text_items",
    ) + SELECTABLE_RETURN_NAMES
    OUTPUT_IS_LIST = (False, False, False, False, False, False, False, False, False, False, True, True) + SELECTABLE_OUTPUT_IS_LIST
    FUNCTION = "ingest_text_file"
    CATEGORY = "ZMongo/03 Documents"
    OUTPUT_NODE = True

    def _nonfatal(
        self,
        *,
        payload: dict[str, Any],
        document_name: str = "",
        filename: str = "",
        refresh: str = "",
    ):
        payload = _ensure_payload_dict(payload)
        return (
            _json_text(payload),
            "",
            "",
            document_name or "",
            filename or "",
            "",
            0,
            0,
            False,
            refresh,
            _as_comfy_list([]),
            _as_comfy_list([]),
            *_selectable_tail([]),
        )

    def ingest_text_file(
        self,
        session: Any,
        file_path: str,
        document_name: str,
        collection_name: str,
        field_root: str,
        overwrite: bool,
        metadata_json: str,
        tags_json: str,
        file_path_link: Any = "",
        refresh_token: str = "",
    ):
        primary_path = _clean_scalar(file_path)
        linked_path = _coerce_file_path_link(file_path_link)
        effective_file_path = linked_path or primary_path
        refresh = _dirty_token("document_ingest_text_file", refresh_token, effective_file_path, document_name, collection_name)

        if not effective_file_path:
            return self._nonfatal(
                payload={
                    "success": False,
                    "message": "Text document ingestion skipped. No file path was provided.",
                    "data": {
                        "file_path": primary_path,
                        "file_path_link": _clean_scalar(file_path_link),
                        "checks": [
                            "Connect 06 Document File Browser.file_path_link to this node's file_path_link.",
                            "Or enter a full path to a .pdf, .docx, .txt, .md, .csv, .json, or .log file.",
                        ],
                        "refresh": refresh,
                    },
                    "error": {"type": "MissingFilePath", "msg": "No file path supplied."},
                    "status_code": 400,
                },
                document_name=document_name,
                refresh=refresh,
            )

        path = Path(os.path.expanduser(effective_file_path)).resolve()
        filename = path.name

        if not path.exists() or not path.is_file():
            return self._nonfatal(
                payload={
                    "success": False,
                    "message": f"Text document ingestion skipped. File was not found: {path}",
                    "data": {
                        "effective_file_path": str(path),
                        "exists": path.exists(),
                        "is_file": path.is_file() if path.exists() else False,
                        "checks": [
                            "The file must exist on the machine running ComfyUI.",
                            "Use 06 Document File Browser to avoid browser/client path confusion.",
                        ],
                        "refresh": refresh,
                    },
                    "error": {"type": "FileNotFound", "msg": "File not found."},
                    "status_code": 404,
                },
                document_name=document_name,
                filename=filename,
                refresh=refresh,
            )

        clean_collection = _clean_scalar(collection_name) or DEFAULT_TEXT_DOCUMENT_COLLECTION
        clean_root = _clean_dot_path_segment(field_root, DEFAULT_TEXT_DOCUMENT_FIELD_ROOT)
        clean_document_name = _clean_scalar(document_name) or path.stem

        try:
            extracted_text, extraction_engine = _extract_text_from_supported_document(path)
        except Exception as exc:
            return self._nonfatal(
                payload={
                    "success": False,
                    "message": f"Text extraction failed: {exc}",
                    "data": {
                        "effective_file_path": str(path),
                        "filename": filename,
                        "extension": path.suffix.lower(),
                        "refresh": refresh,
                    },
                    "error": {"type": exc.__class__.__name__, "msg": str(exc)},
                    "status_code": 0,
                },
                document_name=clean_document_name,
                filename=filename,
                refresh=refresh,
            )

        if not extracted_text.strip():
            return self._nonfatal(
                payload={
                    "success": False,
                    "message": "No text was extracted. This may be a scanned PDF; use OCR ingestion instead.",
                    "data": {
                        "effective_file_path": str(path),
                        "filename": filename,
                        "extension": path.suffix.lower(),
                        "extraction_engine": extraction_engine,
                        "refresh": refresh,
                    },
                    "error": {"type": "NoExtractedText", "msg": "No text extracted."},
                    "status_code": 422,
                },
                document_name=clean_document_name,
                filename=filename,
                refresh=refresh,
            )

        text_size_bytes = _utf8_size(extracted_text)
        if _is_local_storage_session(session) and text_size_bytes > DEFAULT_LOCAL_TEXT_DOCUMENT_MAX_BYTES:
            return self._nonfatal(
                payload=_local_storage_limit_payload(
                    path=path,
                    filename=filename,
                    text_size_bytes=text_size_bytes,
                    refresh=refresh,
                ),
                document_name=clean_document_name,
                filename=filename,
                refresh=refresh,
            )

        now = time.time()
        metadata = _parse_json_object(metadata_json, "metadata_json")
        tags = _parse_json_list(tags_json, "tags_json")

        text_payload = {
            "raw_text": extracted_text,
            "text_length": len(extracted_text),
            "text_size_bytes": text_size_bytes,
            "line_count": len(extracted_text.splitlines()),
            "chunked": False,
            "chunk_count": 0,
            "local_max_text_document_bytes": DEFAULT_LOCAL_TEXT_DOCUMENT_MAX_BYTES,
            "large_file_strategy": "Use hosted ZMongo on BusinessProcessApplications.com; backend ZEmbedder.py handles chunking and embeddings.",
            "extraction_engine": extraction_engine,
        }

        document = {
            "schema_kind": TEXT_DOCUMENT_SCHEMA_KIND,
            "schema_version": TEXT_DOCUMENT_SCHEMA_VERSION,
            "document_name": clean_document_name,
            "collection_name": clean_collection,
            "source_file": {
                "path": str(path),
                "filename": filename,
                "extension": path.suffix.lower(),
                "content_type": mimetypes.guess_type(str(path))[0] or "application/octet-stream",
                "size_bytes": int(path.stat().st_size),
                "ingested_at_unix": now,
            },
            clean_root: text_payload,
            "text": extracted_text,
            "text_length": len(extracted_text),
            "metadata": {
                "tags": tags,
                "user_metadata": metadata,
            },
            "created_at_unix": now,
            "updated_at_unix": now,
        }

        existing_doc, existing_id, existing_query_payload = _query_existing_text_document(
            session,
            collection_name=clean_collection,
            document_name=clean_document_name,
        )

        if existing_id and not bool(overwrite):
            payload = _ensure_payload_dict({
                "success": False,
                "message": f"Text document {clean_document_name!r} already exists. Enable overwrite to replace/update it.",
                "data": {
                    "collection_name": clean_collection,
                    "document_name": clean_document_name,
                    "document_id": existing_id,
                    "existing_query_payload": existing_query_payload,
                    "refresh": refresh,
                },
                "error": {"type": "DocumentExists", "msg": "overwrite is false."},
                "status_code": 409,
            })
            ids = [existing_id]
            return (
                _json_text(payload),
                existing_id,
                existing_id,
                clean_document_name,
                filename,
                extracted_text,
                len(extracted_text),
                DEFAULT_LOCAL_TEXT_DOCUMENT_MAX_BYTES,
                False,
                refresh,
                _as_comfy_list(ids),
                _as_comfy_list([extracted_text]),
                *_selectable_tail(ids),
            )

        if existing_id:
            payload = _save_text_document_by_dot_paths(
                session,
                collection_name=clean_collection,
                document_id=existing_id,
                document=document,
                field_root=clean_root,
            )
            document_id = existing_id
        else:
            payload = _generic_document_create_fallback(
                session,
                document=document,
                collection_name=clean_collection,
            )
            document_id = _extract_doc_id(payload)

        success = bool(payload.get("success"))
        ids = [document_id] if document_id else []
        response_payload = _ensure_payload_dict({
            "success": success,
            "message": (
                f"Ingested {filename!r} into collection {clean_collection!r}."
                if success else
                f"Text was extracted from {filename!r}, but database save failed."
            ),
            "data": {
                "document_id": document_id,
                "document_name": clean_document_name,
                "collection_name": clean_collection,
                "field_root": clean_root,
                "filename": filename,
                "text_length": len(extracted_text),
                "chunked": False,
                "chunk_count": 0,
                "local_max_text_document_bytes": DEFAULT_LOCAL_TEXT_DOCUMENT_MAX_BYTES,
                "save_payload": payload,
                "refresh": refresh,
            },
            "error": None if success else payload.get("error"),
            "status_code": int(payload.get("status_code") or (200 if success else 0)),
        })

        return (
            _json_text(response_payload),
            document_id,
            document_id,
            clean_document_name,
            filename,
            extracted_text,
            len(extracted_text),
            DEFAULT_LOCAL_TEXT_DOCUMENT_MAX_BYTES,
            success,
            refresh,
            _as_comfy_list(ids),
            _as_comfy_list([extracted_text] if extracted_text else []),
            *_selectable_tail(ids),
        )


class ZMongoDocumentCreateText(AlwaysDirtyMixin):
    @classmethod
    def INPUT_TYPES(cls):
        return {"required": {"session": ("ZMONGO_API_SESSION",), "filename": ("STRING", {"default": "manual_document.txt"}), "text": ("STRING", {"default": "", "multiline": True}), "case_id": ("STRING", {"default": ""}), "metadata_json": ("STRING", {"default": "{}", "multiline": True}), "document_json": ("STRING", {"default": "{}", "multiline": True}), "document_prefix": ("STRING", {"default": DEFAULT_DOCUMENT_PREFIX})}, "optional": {"refresh_token": ("STRING", {"default": ""})}}

    RETURN_TYPES = ("STRING", "STRING", ZMONGO_DOCUMENT_ID, "STRING", ZMONGO_FILENAME, "INT", "BOOLEAN", "STRING", "*", "*") + SELECTABLE_RETURN_TYPES
    RETURN_NAMES = ("json", "document_id", "document_id_link", "filename", "filename_link", "text_length", "success", "refresh", "document_ids", "filenames") + SELECTABLE_RETURN_NAMES
    OUTPUT_IS_LIST = (False, False, False, False, False, False, False, False, True, True) + SELECTABLE_OUTPUT_IS_LIST
    FUNCTION = "create_text_document"
    CATEGORY = "ZMongo/03 Documents"
    OUTPUT_NODE = True

    def create_text_document(self, session: Any, filename: str, text: str, case_id: str, metadata_json: str, document_json: str, document_prefix: str, refresh_token: str = ""):
        body = {"filename": _clean_scalar(filename) or "manual_document.txt", "text": text or "", "case_id": _clean_scalar(case_id) or None, "metadata": _parse_json_object(metadata_json, "metadata_json"), "document": _parse_json_object(document_json, "document_json")}
        payload = _session_request(session, "POST", document_prefix, "/api/create", json_body=body)
        if not payload.get("success") and _is_route_not_found(payload):
            document = dict(body.get("document") or {})
            document.update({
                "filename": body["filename"],
                "text": body["text"],
                "case_id": body.get("case_id"),
                "metadata": body.get("metadata") or {},
                "content_type": "text/plain",
                "text_length": len(body["text"]),
            })
            payload = _generic_document_create_fallback(session, document=document)
        _raise_if_failed(payload, "Document create")
        data = _extract_data(payload)
        document_id = _extract_doc_id(payload)
        out_filename = str(data.get("filename") or body["filename"])
        ids = [document_id] if document_id else []
        filenames = [out_filename] if out_filename else []
        return (_json_text(payload), document_id, document_id, out_filename, out_filename, int(data.get("text_length") or len(text or "")), bool(payload.get("success")), _dirty_token("document_create", refresh_token), _as_comfy_list(ids), _as_comfy_list(filenames), *_selectable_tail(ids))


class ZMongoDocumentList(AlwaysDirtyMixin):
    @classmethod
    def INPUT_TYPES(cls):
        return {"required": {"session": ("ZMONGO_API_SESSION",), "query_json": ("STRING", {"default": "{}", "multiline": True}), "sort_json": ("STRING", {"default": '[["updated_at", -1]]', "multiline": True}), "limit": ("INT", {"default": 50, "min": 1, "max": 500}), "skip": ("INT", {"default": 0, "min": 0, "max": 1000000}), "include_file_data": ("BOOLEAN", {"default": False}), "document_prefix": ("STRING", {"default": DEFAULT_DOCUMENT_PREFIX})}, "optional": {"refresh_token": ("STRING", {"default": ""})}}

    RETURN_TYPES = ("STRING", "STRING", "STRING", "STRING", ZMONGO_DOCUMENT_ID, ZMONGO_FILENAME, "INT", "BOOLEAN", "STRING", "*", "*") + SELECTABLE_RETURN_TYPES
    RETURN_NAMES = ("json", "documents_json", "document_ids_json", "filenames_json", "first_document_id_link", "first_filename_link", "count", "success", "refresh", "document_ids", "filenames") + SELECTABLE_RETURN_NAMES
    OUTPUT_IS_LIST = (False, False, False, False, False, False, False, False, False, True, True) + SELECTABLE_OUTPUT_IS_LIST
    FUNCTION = "list_documents"
    CATEGORY = "ZMongo/03 Documents"
    OUTPUT_NODE = True

    def list_documents(self, session: Any, query_json: str, sort_json: str, limit: int, skip: int, include_file_data: bool, document_prefix: str, refresh_token: str = ""):
        body = {"query": _parse_json_object(query_json, "query_json"), "sort": _parse_json_list(sort_json, "sort_json"), "limit": int(limit), "skip": int(skip), "include_file_data": bool(include_file_data)}
        payload = _session_request(session, "POST", document_prefix, "/api/list", json_body=body)
        if not payload.get("success") and _is_route_not_found(payload):
            payload = _generic_document_query_fallback(
                session,
                query=body["query"],
                sort=body["sort"],
                limit=int(limit),
                skip=int(skip),
            )
        _raise_if_failed(payload, "Document list")
        docs = _extract_documents(payload)
        ids = _extract_doc_ids_from_documents(docs)
        filenames = _extract_filenames_from_documents(docs)
        return (_json_text(payload), _json_text(docs), _json_text(ids), _json_text(filenames), (ids[0] if ids else ""), (filenames[0] if filenames else ""), len(docs), bool(payload.get("success")), _dirty_token("document_list", refresh_token), _as_comfy_list(ids), _as_comfy_list(filenames), *_selectable_tail(ids))

class ZMongoDocumentGet(AlwaysDirtyMixin):
    @classmethod
    def INPUT_TYPES(cls):
        return {"required": {"session": ("ZMONGO_API_SESSION",), "document_id": ("STRING", {"default": ""}), "include_file_data": ("BOOLEAN", {"default": False}), "document_prefix": ("STRING", {"default": DEFAULT_DOCUMENT_PREFIX})}, "optional": {"document_id_link": ("*",), "refresh_token": ("STRING", {"default": ""})}}

    RETURN_TYPES = ("STRING", "STRING", "STRING", "STRING", "STRING", "BOOLEAN", "STRING", "*") + SELECTABLE_RETURN_TYPES
    RETURN_NAMES = ("json", "document_id", "document_json", "text", "summary", "success", "refresh", "document_ids") + SELECTABLE_RETURN_NAMES
    OUTPUT_IS_LIST = (False, False, False, False, False, False, False, True) + SELECTABLE_OUTPUT_IS_LIST
    FUNCTION = "get_document"
    CATEGORY = "ZMongo/03 Documents"
    OUTPUT_NODE = True

    def get_document(self, session: Any, document_id: str, include_file_data: bool, document_prefix: str, document_id_link: str = "", refresh_token: str = ""):
        clean_id = _coerce_document_id(_prefer_link_value(document_id, document_id_link))
        quoted = urllib.parse.quote(clean_id, safe="")
        payload = _session_request(session, "GET", document_prefix, f"/api/get/{quoted}", params={"include_file_data": "true" if include_file_data else "false"})
        if not payload.get("success") and _is_route_not_found(payload):
            payload = _generic_document_get_fallback(session, clean_id)
        _raise_if_failed(payload, "Document get")
        document = _extract_document(payload)
        resolved_id = _coerce_document_id(document) or clean_id
        ids = [resolved_id] if resolved_id else []
        return (_json_text(payload), resolved_id, _json_text(document), str(document.get("text") or ""), _document_summary(document), bool(payload.get("success")), _dirty_token("document_get", refresh_token), _as_comfy_list(ids), *_selectable_tail(ids))


class ZMongoDocumentGetText(AlwaysDirtyMixin):
    @classmethod
    def INPUT_TYPES(cls):
        return {
            "required": {
                "session": ("ZMONGO_API_SESSION",),
                "document_id": ("STRING", {"default": ""}),
                "document_prefix": ("STRING", {"default": DEFAULT_DOCUMENT_PREFIX}),
            },
            "optional": {
                "document_id_link": ("*",),
                "collection_name": ("STRING", {"default": DEFAULT_DOCUMENT_COLLECTION}),
                "collection_name_link": ("*",),
                "fallback_text": ("STRING", {"default": "", "multiline": True}),
                "refresh_token": ("STRING", {"default": ""}),
            },
        }

    RETURN_TYPES = ("STRING", "STRING", "INT", "BOOLEAN", "STRING", "*") + SELECTABLE_RETURN_TYPES
    RETURN_NAMES = ("json", "text", "text_length", "success", "refresh", "text_items") + SELECTABLE_RETURN_NAMES
    OUTPUT_IS_LIST = (False, False, False, False, False, True) + SELECTABLE_OUTPUT_IS_LIST
    FUNCTION = "get_text"
    CATEGORY = "ZMongo/03 Documents"
    OUTPUT_NODE = True

    def _nonfatal_get_text_result(
        self,
        *,
        payload: dict[str, Any],
        text: str = "",
        refresh: str,
    ):
        text = text or ""
        items = [text] if text else []
        return (
            _json_text(_ensure_payload_dict(payload)),
            text,
            len(text),
            bool(payload.get("success")),
            refresh,
            _as_comfy_list(items),
            *_selectable_tail(items),
        )

    def get_text(
        self,
        session: Any,
        document_id: str,
        document_prefix: str,
        document_id_link: str = "",
        collection_name: str = DEFAULT_DOCUMENT_COLLECTION,
        collection_name_link: Any = "",
        fallback_text: str = "",
        refresh_token: str = "",
    ):
        clean_id = _coerce_document_id(_prefer_link_value(document_id, document_id_link))
        clean_collection = _clean_scalar(_prefer_link_value(collection_name, collection_name_link)) or DEFAULT_DOCUMENT_COLLECTION
        refresh = _dirty_token("document_text", refresh_token, clean_collection, clean_id)

        if not clean_id:
            payload = _ensure_payload_dict({
                "success": False,
                "message": "Document text get skipped: document_id is empty.",
                "data": {
                    "collection_name": clean_collection,
                    "document_id": clean_id,
                    "refresh": refresh,
                    "checks": [
                        "Connect document_id_link from 06 Select Nth Document Item.",
                        "Confirm the selected value is a document id, not a filename, file path, or field path.",
                    ],
                },
                "error": {"type": "MissingDocumentId", "msg": "document_id is required."},
                "status_code": 400,
            })
            return self._nonfatal_get_text_result(payload=payload, refresh=refresh)

        quoted = urllib.parse.quote(clean_id, safe="")
        text_payload = _session_request(session, "GET", document_prefix, f"/api/text/{quoted}")
        text = _extract_text(text_payload)

        if text_payload.get("success") or text:
            if not text_payload.get("success") and text:
                text_payload = _ensure_payload_dict({
                    "success": True,
                    "message": "Document text route returned a non-success payload, but text was extracted from the response.",
                    "data": {
                        "collection_name": clean_collection,
                        "document_id": clean_id,
                        "text": text,
                        "text_length": len(text),
                        "source_payload": text_payload,
                        "refresh": refresh,
                    },
                    "error": None,
                    "status_code": 200,
                })
            return self._nonfatal_get_text_result(payload=text_payload, text=text, refresh=refresh)

        # Route-based document APIs may not exist for local_file_store sessions.
        # Fall back to the generic local/session document loader and extract text
        # from common fields.  This also fixes documents saved outside the default
        # "documents" collection by allowing collection_name to be supplied.
        fallback_payload = _generic_document_get_fallback(
            session,
            clean_id,
            collection_name=clean_collection,
        )
        fallback_document = _extract_document(fallback_payload)
        fallback_extracted_text = _extract_text(fallback_payload)
        best_document_text, best_document_text_path, best_document_text_score = _best_text_from_document(fallback_document)

        if best_document_text and (
            not fallback_extracted_text
            or _looks_like_bad_text_value(fallback_extracted_text)
            or _text_quality_score(best_document_text) > _text_quality_score(fallback_extracted_text)
        ):
            fallback_extracted_text = best_document_text

        if fallback_payload.get("success") or fallback_extracted_text:
            if fallback_extracted_text:
                combined_payload = _ensure_payload_dict({
                    "success": True,
                    "message": f"Loaded document text using generic/local document fallback from {best_document_text_path or 'payload text'}.",
                    "data": {
                        "collection_name": clean_collection,
                        "document_id": clean_id,
                        "text": fallback_extracted_text,
                        "text_length": len(fallback_extracted_text),
                        "resolved_text_path": best_document_text_path or "payload",
                        "text_quality_score": best_document_text_score,
                        "source": "generic_document_get_fallback",
                        "text_route_payload": text_payload,
                        "fallback_payload": fallback_payload,
                        "refresh": refresh,
                    },
                    "error": None,
                    "status_code": 200,
                })
                return self._nonfatal_get_text_result(
                    payload=combined_payload,
                    text=fallback_extracted_text,
                    refresh=refresh,
                )

        # Final behavior is intentionally nonfatal.  A stale local id or wrong
        # collection should not crash the whole ComfyUI graph.
        final_text = fallback_text or ""
        final_payload = _ensure_payload_dict({
            "success": bool(final_text),
            "message": (
                "Document text was not found; returned fallback_text."
                if final_text else
                "Document text was not found. The workflow continued without raising."
            ),
            "data": {
                "collection_name": clean_collection,
                "document_id": clean_id,
                "text": final_text,
                "text_length": len(final_text),
                "text_route_payload": text_payload,
                "fallback_payload": fallback_payload,
                "refresh": refresh,
                "checks": [
                    "The selected local document id may be stale; run 06 List Documents again and select a current id.",
                    "Confirm collection_name matches the collection used when the document was created or ingested.",
                    "For text files ingested by 06 Ingest Text PDF / DOCX, try collection_name='documents' unless you changed it.",
                    "Confirm the local_store JSON path shown in the payload exists on the ComfyUI machine.",
                ],
            },
            "error": {
                "type": "DocumentTextNotFound",
                "msg": "Document text could not be loaded from the route or generic fallback.",
            },
            "status_code": 404,
        })
        return self._nonfatal_get_text_result(payload=final_payload, text=final_text, refresh=refresh)


class ZMongoDocumentFieldPaths(AlwaysDirtyMixin):
    @classmethod
    def INPUT_TYPES(cls):
        return {
            "required": {
                "session": ("ZMONGO_API_SESSION",),
                "collection_name": ("STRING", {"default": "documents"}),
                "document_id": ("STRING", {"default": ""}),
                "document_prefix": ("STRING", {"default": DEFAULT_DOCUMENT_PREFIX}),
            },
            "optional": {
                "document_id_link": ("*",),
                "collection_name_link": ("*",),
                "refresh_token": ("STRING", {"default": ""}),
            },
        }

    RETURN_TYPES = ("STRING", "STRING", "STRING", "INT", "BOOLEAN", "STRING", "*") + SELECTABLE_RETURN_TYPES
    RETURN_NAMES = ("json", "field_paths_json", "indexed_field_paths", "count", "success", "refresh", "field_paths") + SELECTABLE_RETURN_NAMES
    OUTPUT_IS_LIST = (False, False, False, False, False, False, True) + SELECTABLE_OUTPUT_IS_LIST
    FUNCTION = "field_paths"
    CATEGORY = "ZMongo/03 Documents"
    OUTPUT_NODE = True

    def _local_field_paths_fallback(
        self,
        session: Any,
        *,
        collection_name: str,
        document_id: str,
        refresh: str,
        prior_payload: Optional[dict[str, Any]] = None,
    ):
        local_payload = _generic_document_get_fallback(
            session,
            document_id,
            collection_name=collection_name or DEFAULT_DOCUMENT_COLLECTION,
        )
        document = _extract_document(local_payload)
        if not document:
            payload = _ensure_payload_dict({
                "success": False,
                "message": "Document field paths failed. Local fallback could not load the document.",
                "data": {
                    "collection_name": collection_name,
                    "document_id": document_id,
                    "prior_payload": prior_payload,
                    "local_payload": local_payload,
                    "refresh": refresh,
                    "checks": [
                        "Confirm collection_name matches the Local File Store collection.",
                        "Confirm document_id exists in that local collection.",
                        "Use 03 List Docs or 06 List Documents to verify the document id.",
                    ],
                },
                "error": {"msg": "Document not found in local fallback."},
                "status_code": local_payload.get("status_code", 404),
            })
            return (_json_text(payload), _json_text([]), _indexed_list_text([]), 0, False, refresh, _as_comfy_list([]), *_selectable_tail([]))

        paths = _dedupe_strings(_flatten_document_paths(document))
        payload = _ensure_payload_dict({
            "success": True,
            "message": f"Found {len(paths)} field path(s) using local session fallback.",
            "data": {
                "collection_name": collection_name,
                "document_id": document_id,
                "field_paths": paths,
                "count": len(paths),
                "source": "local_session_get_doc_fallback",
                "prior_payload": prior_payload,
                "refresh": refresh,
            },
            "error": None,
            "status_code": 200,
        })
        return (_json_text(payload), _json_text(paths), _indexed_list_text(paths), len(paths), True, refresh, _as_comfy_list(paths), *_selectable_tail(paths))

    def field_paths(
        self,
        session: Any,
        collection_name: str,
        document_id: str,
        document_prefix: str,
        document_id_link: Any = "",
        collection_name_link: Any = "",
        refresh_token: str = "",
    ):
        clean_id = _coerce_document_id(_prefer_link_value(document_id, document_id_link))
        clean_collection = _clean_scalar(_prefer_link_value(collection_name, collection_name_link)) or "documents"
        refresh = _dirty_token("document_field_paths", refresh_token, clean_collection, clean_id)

        if not clean_id:
            payload = _ensure_payload_dict({
                "success": False,
                "message": "Document field paths skipped: document_id is empty.",
                "data": {
                    "collection_name": clean_collection,
                    "document_id": clean_id,
                    "refresh": refresh,
                    "checks": [
                        "Connect document_id_link from 06 Select Nth Document Item.",
                        "For Local File Store, select an actual document id, not a field path.",
                    ],
                },
                "error": {"msg": "missing_document_id"},
                "status_code": 400,
            })
            return (_json_text(payload), _json_text([]), _indexed_list_text([]), 0, False, refresh, _as_comfy_list([]), *_selectable_tail([]))

        quoted = urllib.parse.quote(clean_id, safe="")
        payload = _session_request(session, "GET", document_prefix, f"/api/field-paths/{quoted}")

        if payload.get("success"):
            paths = _extract_field_paths(payload)
            return (_json_text(payload), _json_text(paths), _indexed_list_text(paths), len(paths), True, refresh, _as_comfy_list(paths), *_selectable_tail(paths))

        # Local File Store and generic ZMongo sessions do not implement the
        # /documents field-path route.  Fall back to get_doc(collection, id) and
        # flatten the document locally.
        return self._local_field_paths_fallback(
            session,
            collection_name=clean_collection,
            document_id=clean_id,
            refresh=refresh,
            prior_payload=payload,
        )

class ZMongoDocumentGetValue(AlwaysDirtyMixin):
    @classmethod
    def INPUT_TYPES(cls):
        return {"required": {"session": ("ZMONGO_API_SESSION",), "document_id": ("STRING", {"default": ""}), "field_path": ("STRING", {"default": "text"}), "fallback": ("STRING", {"default": ""}), "document_prefix": ("STRING", {"default": DEFAULT_DOCUMENT_PREFIX})}, "optional": {"document_id_link": ("*",), "field_path_link": ("*",), "refresh_token": ("STRING", {"default": ""})}}

    RETURN_TYPES = ("STRING", "STRING", "BOOLEAN", "STRING", "STRING", "*") + SELECTABLE_RETURN_TYPES
    RETURN_NAMES = ("json", "value", "exists", "value_type", "refresh", "value_items") + SELECTABLE_RETURN_NAMES
    OUTPUT_IS_LIST = (False, False, False, False, False, True) + SELECTABLE_OUTPUT_IS_LIST
    FUNCTION = "get_value"
    CATEGORY = "ZMongo/03 Documents"
    OUTPUT_NODE = True

    def get_value(self, session: Any, document_id: str, field_path: str, fallback: str, document_prefix: str, document_id_link: str = "", field_path_link: str = "", refresh_token: str = ""):
        clean_id = _coerce_document_id(_prefer_link_value(document_id, document_id_link))
        clean_path = _coerce_field_path(_prefer_link_value(field_path, field_path_link))
        refresh = _dirty_token("document_get_value", refresh_token, clean_id, clean_path)
        quoted = urllib.parse.quote(clean_id, safe="")

        payload = _session_request(session, "GET", document_prefix, f"/api/get/{quoted}")
        if not payload.get("success") and _is_route_not_found(payload):
            payload = _generic_document_get_fallback(session, clean_id)

        if not payload.get("success"):
            empty_tail = _selectable_tail([])
            return (_json_text(payload), fallback or "", False, "missing_document", _dirty_token("document_get_value_error", refresh_token), _as_comfy_list([]), *empty_tail)

        document = _extract_document(payload)
        marker = object()
        value = _safe_get_by_path(document, clean_path, marker)
        exists = value is not marker

        resolved_path = clean_path
        smart_replacement_used = False
        if not exists:
            value = fallback
        elif clean_path in {"text", "raw_text", "document_text.raw_text", "document_text.text"} and _looks_like_bad_text_value(value):
            best_text, best_path, best_score = _best_text_from_document(document)
            if best_text and best_path and best_path != clean_path and _text_quality_score(best_text) > _text_quality_score(value):
                value = best_text
                resolved_path = best_path
                smart_replacement_used = True

        value_text = _json_text(value) if isinstance(value, (dict, list, tuple)) else str(value if value is not None else "")
        items = _value_items(value)

        if smart_replacement_used:
            payload = _ensure_payload_dict({
                "success": True,
                "message": f"Field path {clean_path!r} looked binary/corrupt; returned better text from {resolved_path!r}.",
                "data": {
                    "document_id": clean_id,
                    "requested_field_path": clean_path,
                    "resolved_path": resolved_path,
                    "exists": True,
                    "value_type": type(value).__name__,
                    "value": value,
                    "value_text": value_text,
                    "refresh": refresh,
                },
                "error": None,
                "status_code": 200,
            })

        return (_json_text(payload), value_text, bool(exists), type(value).__name__ if exists else "missing", refresh, _as_comfy_list(items), *_selectable_tail(items))


class ZMongoDocumentSaveText(AlwaysDirtyMixin):
    @classmethod
    def INPUT_TYPES(cls):
        return {"required": {"session": ("ZMONGO_API_SESSION",), "document_id": ("STRING", {"default": ""}), "text": ("STRING", {"default": "", "multiline": True}), "document_prefix": ("STRING", {"default": DEFAULT_DOCUMENT_PREFIX})}, "optional": {"document_id_link": ("*",), "refresh_token": ("STRING", {"default": ""})}}

    RETURN_TYPES = ("STRING", "STRING", "INT", "BOOLEAN", "STRING", "*") + SELECTABLE_RETURN_TYPES
    RETURN_NAMES = ("json", "document_id", "text_length", "success", "refresh", "document_ids") + SELECTABLE_RETURN_NAMES
    OUTPUT_IS_LIST = (False, False, False, False, False, True) + SELECTABLE_OUTPUT_IS_LIST
    FUNCTION = "save_text"
    CATEGORY = "ZMongo/03 Documents"
    OUTPUT_NODE = True

    def save_text(self, session: Any, document_id: str, text: str, document_prefix: str, document_id_link: str = "", refresh_token: str = ""):
        clean_id = _coerce_document_id(_prefer_link_value(document_id, document_id_link))
        quoted = urllib.parse.quote(clean_id, safe="")
        payload = _session_request(session, "POST", document_prefix, f"/api/save-text/{quoted}", json_body={"text": text or ""})
        if not payload.get("success") and _is_route_not_found(payload):
            payload = _session_save_value(
                session,
                collection_name=DEFAULT_DOCUMENT_COLLECTION,
                document_id=clean_id,
                query=None,
                field_path="text",
                value=text or "",
                upsert=False,
            )
        _raise_if_failed(payload, "Document save text")
        resolved_id = _extract_doc_id(payload) or clean_id
        ids = [resolved_id] if resolved_id else []
        return (_json_text(payload), resolved_id, int(_extract_data(payload).get("text_length") or len(text or "")), bool(payload.get("success")), _dirty_token("document_save_text", refresh_token), _as_comfy_list(ids), *_selectable_tail(ids))


class ZMongoDocumentSaveValue(AlwaysDirtyMixin):
    @classmethod
    def INPUT_TYPES(cls):
        return {"required": {"session": ("ZMONGO_API_SESSION",), "document_id": ("STRING", {"default": ""}), "field_path": ("STRING", {"default": "metadata.note"}), "value": ("STRING", {"default": "", "multiline": True}), "parse_json": ("BOOLEAN", {"default": True}), "normalize_for_storage": ("BOOLEAN", {"default": False}), "document_prefix": ("STRING", {"default": DEFAULT_DOCUMENT_PREFIX})}, "optional": {"document_id_link": ("*",), "field_path_link": ("*",), "refresh_token": ("STRING", {"default": ""})}}

    RETURN_TYPES = ("STRING", "STRING", "BOOLEAN", "STRING", "*") + SELECTABLE_RETURN_TYPES
    RETURN_NAMES = ("json", "document_id", "success", "refresh", "document_ids") + SELECTABLE_RETURN_NAMES
    OUTPUT_IS_LIST = (False, False, False, False, True) + SELECTABLE_OUTPUT_IS_LIST
    FUNCTION = "save_value"
    CATEGORY = "ZMongo/03 Documents"
    OUTPUT_NODE = True

    def save_value(self, session: Any, document_id: str, field_path: str, value: str, parse_json: bool, normalize_for_storage: bool, document_prefix: str, document_id_link: str = "", field_path_link: str = "", refresh_token: str = ""):
        clean_id = _coerce_document_id(_prefer_link_value(document_id, document_id_link))
        clean_path = _coerce_field_path(_prefer_link_value(field_path, field_path_link))
        refresh = _dirty_token("document_save_value", refresh_token, clean_id, clean_path)
        ids = [clean_id] if clean_id else []

        # MongoDB will reject attempts to update _id.  More importantly, these
        # fields are identity/ownership fields and should never be written by a
        # generic ComfyUI Save Value node.  Return a visible, non-throwing error
        # payload so the workflow continues and the user can choose a writable
        # field path from the selector.
        if _is_immutable_document_field_path(clean_path):
            payload = _blocked_immutable_field_payload(
                document_id=clean_id,
                field_path=clean_path,
                refresh=refresh,
            )
            return (_json_text(payload), clean_id, False, refresh, _as_comfy_list(ids), *_selectable_tail(ids))

        if not clean_path:
            payload = _ensure_payload_dict(
                {
                    "success": False,
                    "message": "Save skipped. field_path is required.",
                    "data": {"document_id": clean_id, "field_path": clean_path, "refresh": refresh},
                    "error": {"type": "MissingFieldPath", "msg": "field_path is required."},
                    "status_code": 400,
                }
            )
            return (_json_text(payload), clean_id, False, refresh, _as_comfy_list(ids), *_selectable_tail(ids))

        quoted = urllib.parse.quote(clean_id, safe="")
        body = {
            "field_path": clean_path,
            "value": _parse_any_json(value, parse_json=parse_json),
            "parse_json_strings": bool(parse_json),
            "normalize_for_storage": bool(normalize_for_storage),
        }
        payload = _session_request(session, "POST", document_prefix, f"/api/save-value/{quoted}", json_body=body)
        if not payload.get("success") and _is_route_not_found(payload):
            payload = _session_save_value(
                session,
                collection_name=DEFAULT_DOCUMENT_COLLECTION,
                document_id=clean_id,
                query=None,
                field_path=clean_path,
                value=body["value"],
                upsert=False,
            )

        # Do not raise here.  Save Value is often used in demo workflows with
        # selectable fields; a failed write should be visible in the JSON output
        # but should not kill the entire ComfyUI graph.
        resolved_id = _extract_doc_id(payload) or clean_id
        ids = [resolved_id] if resolved_id else []
        return (_json_text(payload), resolved_id, bool(payload.get("success")), refresh, _as_comfy_list(ids), *_selectable_tail(ids))


class ZMongoDocumentUpdateMetadata(AlwaysDirtyMixin):
    @classmethod
    def INPUT_TYPES(cls):
        return {"required": {"session": ("ZMONGO_API_SESSION",), "document_id": ("STRING", {"default": ""}), "metadata_json": ("STRING", {"default": "{}", "multiline": True}), "case_id": ("STRING", {"default": ""}), "status": ("STRING", {"default": ""}), "title": ("STRING", {"default": ""}), "description": ("STRING", {"default": "", "multiline": True}), "tags_json": ("STRING", {"default": "[]", "multiline": True}), "document_prefix": ("STRING", {"default": DEFAULT_DOCUMENT_PREFIX})}, "optional": {"document_id_link": ("*",), "refresh_token": ("STRING", {"default": ""})}}

    RETURN_TYPES = ("STRING", "STRING", "BOOLEAN", "STRING", "*") + SELECTABLE_RETURN_TYPES
    RETURN_NAMES = ("json", "document_id", "success", "refresh", "document_ids") + SELECTABLE_RETURN_NAMES
    OUTPUT_IS_LIST = (False, False, False, False, True) + SELECTABLE_OUTPUT_IS_LIST
    FUNCTION = "update_metadata"
    CATEGORY = "ZMongo/03 Documents"
    OUTPUT_NODE = True

    def update_metadata(self, session: Any, document_id: str, metadata_json: str, case_id: str, status: str, title: str, description: str, tags_json: str, document_prefix: str, document_id_link: str = "", refresh_token: str = ""):
        clean_id = _coerce_document_id(_prefer_link_value(document_id, document_id_link))
        quoted = urllib.parse.quote(clean_id, safe="")
        body: dict[str, Any] = {"metadata": _parse_json_object(metadata_json, "metadata_json")}
        if _clean_scalar(case_id):
            body["case_id"] = _clean_scalar(case_id)
        if _clean_scalar(status):
            body["status"] = _clean_scalar(status)
        if _clean_scalar(title):
            body["title"] = _clean_scalar(title)
        if _clean_scalar(description):
            body["description"] = description
        tags = _parse_json_list(tags_json, "tags_json")
        if tags:
            body["tags"] = tags
        payload = _session_request(session, "POST", document_prefix, f"/api/metadata/{quoted}", json_body=body)
        _raise_if_failed(payload, "Document metadata update")
        resolved_id = _extract_doc_id(payload) or clean_id
        ids = [resolved_id] if resolved_id else []
        return (_json_text(payload), resolved_id, bool(payload.get("success")), _dirty_token("document_metadata", refresh_token), _as_comfy_list(ids), *_selectable_tail(ids))


class ZMongoDocumentExtractText(AlwaysDirtyMixin):
    @classmethod
    def INPUT_TYPES(cls):
        return {"required": {"session": ("ZMONGO_API_SESSION",), "document_id": ("STRING", {"default": ""}), "save": ("BOOLEAN", {"default": True}), "document_prefix": ("STRING", {"default": DEFAULT_DOCUMENT_PREFIX})}, "optional": {"document_id_link": ("*",), "refresh_token": ("STRING", {"default": ""})}}

    RETURN_TYPES = ("STRING", "STRING", "INT", "BOOLEAN", "STRING", "*") + SELECTABLE_RETURN_TYPES
    RETURN_NAMES = ("json", "text", "text_length", "success", "refresh", "text_items") + SELECTABLE_RETURN_NAMES
    OUTPUT_IS_LIST = (False, False, False, False, False, True) + SELECTABLE_OUTPUT_IS_LIST
    FUNCTION = "extract_text"
    CATEGORY = "ZMongo/03 Documents"
    OUTPUT_NODE = True

    def extract_text(self, session: Any, document_id: str, save: bool, document_prefix: str, document_id_link: str = "", refresh_token: str = ""):
        clean_id = _coerce_document_id(_prefer_link_value(document_id, document_id_link))
        refresh = _dirty_token("document_extract_text", refresh_token, clean_id)

        if not clean_id:
            payload = _ensure_payload_dict({
                "success": False,
                "message": "Document extract text skipped: document_id is empty.",
                "data": {"document_id": clean_id, "refresh": refresh},
                "error": {"msg": "missing_document_id"},
                "status_code": 400,
            })
            return (_json_text(payload), "", 0, False, refresh, _as_comfy_list([]), *_selectable_tail([]))

        quoted = urllib.parse.quote(clean_id, safe="")
        payload = _session_request(
            session,
            "POST",
            document_prefix,
            f"/api/extract-text/{quoted}",
            json_body={"save": bool(save)},
        )

        # Do not crash ComfyUI when the selected document is text-only or has no
        # stored file payload.  This is a normal state for documents created by
        # "06 Create Text Document".  Fall back to the document's stored text.
        if not payload.get("success"):
            fallback_payload = _session_request(session, "GET", document_prefix, f"/api/text/{quoted}")
            fallback_text = _extract_text(fallback_payload)

            if fallback_text:
                combined_payload = _ensure_payload_dict({
                    "success": True,
                    "message": "No file payload was available for extraction; returned existing stored text instead.",
                    "data": {
                        "document_id": clean_id,
                        "text": fallback_text,
                        "text_length": len(fallback_text),
                        "extract_payload": payload,
                        "fallback_payload": fallback_payload,
                        "source": "stored_text_fallback",
                        "refresh": refresh,
                    },
                    "error": None,
                    "status_code": 200,
                })
                items = [fallback_text]
                return (
                    _json_text(combined_payload),
                    fallback_text,
                    len(fallback_text),
                    True,
                    refresh,
                    _as_comfy_list(items),
                    *_selectable_tail(items),
                )

            # Nonfatal error: report the backend response, but let the workflow continue.
            nonfatal_payload = _ensure_payload_dict({
                "success": False,
                "message": payload.get("message") or "Document text extraction failed without stored text fallback.",
                "data": {
                    "document_id": clean_id,
                    "extract_payload": payload,
                    "fallback_payload": fallback_payload,
                    "source": "extract_failed_no_fallback_text",
                    "refresh": refresh,
                    "checks": [
                        "Use this node only for uploaded file-backed documents when you need PDF/file extraction.",
                        "For documents created with 06 Create Text Document, use 06 Get Document Text instead.",
                        "Confirm that the selected document_id came from 06 List Documents, not from 06 Document Field Paths.",
                    ],
                },
                "error": payload.get("error") or {"msg": payload.get("message") or "extract_failed"},
                "status_code": payload.get("status_code", 0),
            })
            return (_json_text(nonfatal_payload), "", 0, False, refresh, _as_comfy_list([]), *_selectable_tail([]))

        text = _extract_text(payload)
        items = [text] if text else []
        return (_json_text(payload), text, len(text), bool(payload.get("success")), refresh, _as_comfy_list(items), *_selectable_tail(items))


class ZMongoDocumentQueueOCR(AlwaysDirtyMixin):
    @classmethod
    def INPUT_TYPES(cls):
        return {"required": {"session": ("ZMONGO_API_SESSION",), "document_id": ("STRING", {"default": ""}), "priority": ("INT", {"default": 50, "min": 0, "max": 100}), "source": ("STRING", {"default": "comfyui_zmongo_document_node"}), "document_prefix": ("STRING", {"default": DEFAULT_DOCUMENT_PREFIX})}, "optional": {"document_id_link": ("*",), "refresh_token": ("STRING", {"default": ""})}}

    RETURN_TYPES = ("STRING", "STRING", "STRING", "BOOLEAN", "STRING", "*") + SELECTABLE_RETURN_TYPES
    RETURN_NAMES = ("json", "document_id", "status", "success", "refresh", "document_ids") + SELECTABLE_RETURN_NAMES
    OUTPUT_IS_LIST = (False, False, False, False, False, True) + SELECTABLE_OUTPUT_IS_LIST
    FUNCTION = "queue_ocr"
    CATEGORY = "ZMongo/03 Documents"
    OUTPUT_NODE = True

    def queue_ocr(self, session: Any, document_id: str, priority: int, source: str, document_prefix: str, document_id_link: str = "", refresh_token: str = ""):
        clean_id = _coerce_document_id(_prefer_link_value(document_id, document_id_link))
        quoted = urllib.parse.quote(clean_id, safe="")
        payload = _session_request(session, "POST", document_prefix, f"/api/ocr/queue/{quoted}", json_body={"priority": int(priority), "source": _clean_scalar(source) or "comfyui_zmongo_document_node"})
        _raise_if_failed(payload, "Document OCR queue")
        data = _extract_data(payload)
        resolved_id = _coerce_document_id(data.get("document_id")) or clean_id
        status = str(data.get("status") or _safe_get_by_path(data, "ocr.status", ""))
        ids = [resolved_id] if resolved_id else []
        return (_json_text(payload), resolved_id, status, bool(payload.get("success")), _dirty_token("document_queue_ocr", refresh_token), _as_comfy_list(ids), *_selectable_tail(ids))


class ZMongoDocumentOCRStatus(AlwaysDirtyMixin):
    @classmethod
    def INPUT_TYPES(cls):
        return {"required": {"session": ("ZMONGO_API_SESSION",), "document_id": ("STRING", {"default": ""}), "document_prefix": ("STRING", {"default": DEFAULT_DOCUMENT_PREFIX})}, "optional": {"document_id_link": ("*",), "refresh_token": ("STRING", {"default": ""})}}

    RETURN_TYPES = ("STRING", "STRING", "BOOLEAN", "STRING", "BOOLEAN", "STRING", "*") + SELECTABLE_RETURN_TYPES
    RETURN_NAMES = ("json", "status", "has_text", "last_error", "success", "refresh", "status_items") + SELECTABLE_RETURN_NAMES
    OUTPUT_IS_LIST = (False, False, False, False, False, False, True) + SELECTABLE_OUTPUT_IS_LIST
    FUNCTION = "ocr_status"
    CATEGORY = "ZMongo/03 Documents"
    OUTPUT_NODE = True

    def ocr_status(self, session: Any, document_id: str, document_prefix: str, document_id_link: str = "", refresh_token: str = ""):
        clean_id = _coerce_document_id(_prefer_link_value(document_id, document_id_link))
        quoted = urllib.parse.quote(clean_id, safe="")
        payload = _session_request(session, "GET", document_prefix, f"/api/ocr/status/{quoted}")
        _raise_if_failed(payload, "Document OCR status")
        data = _extract_data(payload)
        status = str(data.get("status") or "")
        last_error = str(data.get("last_error") or "")
        items = [status] if status else []
        return (_json_text(payload), status, bool(data.get("has_text")), last_error, bool(payload.get("success")), _dirty_token("document_ocr_status", refresh_token), _as_comfy_list(items), *_selectable_tail(items))


class ZMongoDocumentDelete(AlwaysDirtyMixin):
    @classmethod
    def INPUT_TYPES(cls):
        return {"required": {"session": ("ZMONGO_API_SESSION",), "document_id": ("STRING", {"default": ""}), "confirm_delete": ("BOOLEAN", {"default": False}), "document_prefix": ("STRING", {"default": DEFAULT_DOCUMENT_PREFIX})}, "optional": {"document_id_link": ("*",), "refresh_token": ("STRING", {"default": ""})}}

    RETURN_TYPES = ("STRING", "STRING", "BOOLEAN", "STRING", "*") + SELECTABLE_RETURN_TYPES
    RETURN_NAMES = ("json", "document_id", "success", "refresh", "document_ids") + SELECTABLE_RETURN_NAMES
    OUTPUT_IS_LIST = (False, False, False, False, True) + SELECTABLE_OUTPUT_IS_LIST
    FUNCTION = "delete_document"
    CATEGORY = "ZMongo/03 Documents"
    OUTPUT_NODE = True

    def delete_document(self, session: Any, document_id: str, confirm_delete: bool, document_prefix: str, document_id_link: str = "", refresh_token: str = ""):
        clean_id = _coerce_document_id(_prefer_link_value(document_id, document_id_link))
        if not confirm_delete:
            payload = _error_payload("Delete not executed. Set confirm_delete=true.", data={"document_id": clean_id}, status_code=400)
            ids = [clean_id] if clean_id else []
            return (_json_text(payload), clean_id, False, _dirty_token("document_delete_blocked", refresh_token), _as_comfy_list(ids), *_selectable_tail(ids))
        quoted = urllib.parse.quote(clean_id, safe="")
        payload = _session_request(session, "POST", document_prefix, f"/api/delete/{quoted}", json_body={})
        _raise_if_failed(payload, "Document delete")
        ids = [clean_id] if clean_id else []
        return (_json_text(payload), clean_id, bool(payload.get("success")), _dirty_token("document_delete", refresh_token), _as_comfy_list(ids), *_selectable_tail(ids))


class ZMongoDocumentSelectNthItem(AlwaysDirtyMixin):
    """
    Color-coded selector for document workflows.

    The selected item is returned as a generic string plus guarded semantic
    outputs. The guarded outputs prevent a selected field path like "_id" from
    being sent as a document id.
    """

    @classmethod
    def INPUT_TYPES(cls):
        return {
            "required": {
                "items_list": ("*",),
                "index": ("INT", {"default": 0, "min": 0, "max": 1000000}),
                "fallback": ("STRING", {"default": ""}),
                "semantic_hint": (
                    ["auto", "document_id", "field_path", "file_path", "filename", "text"],
                    {"default": "auto"},
                ),
            }
        }

    RETURN_TYPES = ("STRING", ZMONGO_DOCUMENT_ID, ZMONGO_FIELD_PATH, ZMONGO_FILE_PATH, ZMONGO_FILENAME, ZMONGO_TEXT, "STRING")
    RETURN_NAMES = ("item", "document_id_link", "field_path_link", "file_path_link", "filename_link", "text_link", "status")
    FUNCTION = "select_nth_item"
    CATEGORY = "ZMongo/03 Documents"
    INPUT_IS_LIST = True
    OUTPUT_NODE = True

    @staticmethod
    def _unwrap_scalar(value: Any, default: Any = None) -> Any:
        if isinstance(value, list):
            if not value:
                return default
            return value[0]
        return value if value is not None else default

    @staticmethod
    def _normalize_items(value: Any) -> list[Any]:
        if value is None:
            return []
        if isinstance(value, tuple):
            value = list(value)
        if isinstance(value, list):
            if len(value) == 1 and isinstance(value[0], (list, tuple)):
                return list(value[0])
            return value
        return [value]

    @staticmethod
    def _selected_as_string(value: Any) -> str:
        if isinstance(value, (dict, list, tuple)):
            return _json_text(value)
        return _strip_index_prefix(value)

    def _semantic_outputs(self, selected: str, semantic_hint: str = "auto") -> tuple[str, str, str, str, str]:
        # ComfyUI sends every input as a list when INPUT_IS_LIST=True.
        # Unwrap semantic_hint before using string methods so the selector can
        # run in both linked/list mode and normal widget mode.
        hint_value = self._unwrap_scalar(semantic_hint, "auto")
        hint = str(hint_value or "auto").strip().lower()

        selected = self._selected_as_string(self._unwrap_scalar(selected, selected))

        document_id = _coerce_document_id_link(selected)
        field_path = _coerce_field_path(selected)
        file_path = _coerce_file_path_link(selected)
        filename = _coerce_filename_link(selected)
        text = _coerce_text_link(selected)

        # Explicit hints produce only the matching semantic output.
        if hint == "document_id":
            return (document_id, "", "", "", "")
        if hint == "field_path":
            return ("", field_path, "", "", "")
        if hint == "file_path":
            return ("", "", file_path, "", "")
        if hint == "filename":
            return ("", "", "", filename, "")
        if hint == "text":
            return ("", "", "", "", selected)

        # Auto mode: never put a simple field-path key into document_id_link.
        if document_id:
            return (document_id, "", "", "", "")
        if file_path:
            return ("", "", file_path, filename, "")
        if field_path:
            return ("", field_path, "", filename, "")
        return ("", "", "", filename, text)

    def select_nth_item(self, items_list, index, fallback, semantic_hint="auto"):
        raw_items = self._normalize_items(items_list)
        fallback_value = str(self._unwrap_scalar(fallback, "") or "")
        index_value = self._unwrap_scalar(index, 0)

        try:
            safe_index = int(index_value or 0)
        except Exception:
            safe_index = 0

        cleaned = [self._selected_as_string(item).strip() for item in raw_items if self._selected_as_string(item).strip()]
        if not cleaned:
            selected = fallback_value
            document_id, field_path, file_path, filename, text = self._semantic_outputs(selected, semantic_hint)
            return (selected, document_id, field_path, file_path, filename, text, "Input list was empty.")

        selected_index = max(0, min(safe_index, len(cleaned) - 1))
        selected = cleaned[selected_index]
        document_id, field_path, file_path, filename, text = self._semantic_outputs(selected, semantic_hint)

        status = (
            f"Selected {selected_index + 1}/{len(cleaned)}: {selected} | "
            f"document_id_link={document_id or '<empty>'} | "
            f"field_path_link={field_path or '<empty>'} | "
            f"file_path_link={file_path or '<empty>'}"
        )
        return (selected, document_id, field_path, file_path, filename, text, status)



class ZMongoDocumentIndexSelector(AlwaysDirtyMixin):
    """
    Select one or more items from a ComfyUI list by single index, range,
    comma-separated series, mixed expression, or all.

    Designed for document workflows:
        06/List Docs.ids -> 06 Document Index Selector -> 06 Get Document Text

    Selection examples:
        3
        2-7
        0,2,5
        0,2-5,9
        *
        all
    """

    @classmethod
    def INPUT_TYPES(cls):
        return {
            "required": {
                "items_list": ("*",),
                "selection": (
                    "STRING",
                    {
                        "default": "0",
                        "tooltip": "Index expression: 0, 2-7, 0,2,5, 0,2-5,9, *, or all.",
                    },
                ),
                "fallback": (
                    "STRING",
                    {
                        "default": "",
                        "tooltip": "Returned only when the input list is empty or no selection matches.",
                    },
                ),
                "clamp_indexes": (
                    "BOOLEAN",
                    {
                        "default": True,
                        "tooltip": "Clamp out-of-range indexes to the nearest valid item instead of dropping them.",
                    },
                ),
                "dedupe": (
                    "BOOLEAN",
                    {
                        "default": True,
                        "tooltip": "Remove duplicate selected items while preserving order.",
                    },
                ),
                "semantic_hint": (
                    ["auto", "document_id", "field_path", "file_path", "filename", "text"],
                    {
                        "default": "document_id",
                        "tooltip": "Controls the guarded semantic link outputs.",
                    },
                ),
            },
            "optional": {
                "refresh_token": ("STRING", {"default": ""}),
            },
        }

    RETURN_TYPES = (
        "*",
        ZMONGO_DOCUMENT_ID,
        "STRING",
        ZMONGO_DOCUMENT_ID,
        ZMONGO_FIELD_PATH,
        ZMONGO_FILE_PATH,
        ZMONGO_FILENAME,
        ZMONGO_TEXT,
        "STRING",
        "STRING",
        "INT",
        "BOOLEAN",
        "STRING",
    )
    RETURN_NAMES = (
        "selected_items",
        "document_id_links",
        "first_item",
        "first_document_id_link",
        "first_field_path_link",
        "first_file_path_link",
        "first_filename_link",
        "first_text_link",
        "selected_json",
        "indexed_selected_items",
        "count",
        "success",
        "status",
    )
    OUTPUT_IS_LIST = (
        True,
        True,
        False,
        False,
        False,
        False,
        False,
        False,
        False,
        False,
        False,
        False,
        False,
    )
    FUNCTION = "select_indexes"
    CATEGORY = "ZMongo/03 Documents"
    INPUT_IS_LIST = True
    OUTPUT_NODE = True

    @staticmethod
    def _unwrap_scalar(value: Any, default: Any = None) -> Any:
        if isinstance(value, list):
            if not value:
                return default
            return value[0]
        return value if value is not None else default

    @staticmethod
    def _normalize_items(value: Any) -> list[Any]:
        if value is None:
            return []
        if isinstance(value, tuple):
            value = list(value)
        if isinstance(value, list):
            if len(value) == 1 and isinstance(value[0], (list, tuple)):
                return list(value[0])
            return value
        return [value]

    @staticmethod
    def _selected_as_string(value: Any) -> str:
        if isinstance(value, (dict, list, tuple)):
            return _json_text(value)
        return _strip_index_prefix(value)

    @staticmethod
    def _parse_index_token(token: str, max_index: int, clamp_indexes: bool) -> list[int]:
        token = str(token or "").strip()
        if not token:
            return []

        # Allow bracket-ish user input copied from previews.
        token = token.strip().strip("[](){}")

        def coerce_index(raw: str) -> Optional[int]:
            try:
                idx = int(str(raw).strip())
            except Exception:
                return None

            if idx < 0:
                idx = max_index + 1 + idx

            if 0 <= idx <= max_index:
                return idx

            if clamp_indexes and max_index >= 0:
                return max(0, min(idx, max_index))

            return None

        # Range syntax: 2-7, 7-2, 2:7, or 2..7. End is inclusive.
        for sep in ("..", ":", "-"):
            if sep in token:
                left, right = token.split(sep, 1)
                start = coerce_index(left)
                end = coerce_index(right)
                if start is None or end is None:
                    return []
                step = 1 if end >= start else -1
                return list(range(start, end + step, step))

        idx = coerce_index(token)
        return [idx] if idx is not None else []

    @classmethod
    def _parse_selection_expression(
        cls,
        selection: str,
        item_count: int,
        clamp_indexes: bool,
    ) -> list[int]:
        max_index = item_count - 1
        expression = str(selection or "").strip().lower()

        if not expression or expression in {"0", "first"}:
            return [0] if item_count else []

        if expression in {"*", "all", "each", "every"}:
            return list(range(item_count))

        if expression in {"last", "-1"}:
            return [max_index] if item_count else []

        # Normalize semicolons and whitespace-separated lists, but preserve range separators.
        expression = expression.replace(";", ",")
        raw_tokens: list[str] = []
        for comma_part in expression.split(","):
            comma_part = comma_part.strip()
            if not comma_part:
                continue
            if " " in comma_part and not any(sep in comma_part for sep in ("-", ":", "..")):
                raw_tokens.extend(part for part in comma_part.split() if part.strip())
            else:
                raw_tokens.append(comma_part)

        indexes: list[int] = []
        for token in raw_tokens:
            indexes.extend(cls._parse_index_token(token, max_index, clamp_indexes))

        return indexes

    @staticmethod
    def _dedupe_indexes(indexes: list[int]) -> list[int]:
        seen: set[int] = set()
        out: list[int] = []
        for idx in indexes:
            if idx in seen:
                continue
            seen.add(idx)
            out.append(idx)
        return out

    def _semantic_outputs(self, selected: str, semantic_hint: str = "document_id") -> tuple[str, str, str, str, str]:
        hint = str(semantic_hint or "document_id").strip().lower()
        selected = self._selected_as_string(selected)

        document_id = _coerce_document_id_link(selected)
        field_path = _coerce_field_path(selected)
        file_path = _coerce_file_path_link(selected)
        filename = _coerce_filename_link(selected)
        text = _coerce_text_link(selected)

        if hint == "document_id":
            return (document_id, "", "", "", "")
        if hint == "field_path":
            return ("", field_path, "", "", "")
        if hint == "file_path":
            return ("", "", file_path, filename, "")
        if hint == "filename":
            return ("", "", "", filename, "")
        if hint == "text":
            return ("", "", "", "", selected)

        if document_id:
            return (document_id, "", "", "", "")
        if file_path:
            return ("", "", file_path, filename, "")
        if field_path:
            return ("", field_path, "", filename, "")
        return ("", "", "", filename, text)

    def select_indexes(
        self,
        items_list,
        selection,
        fallback,
        clamp_indexes,
        dedupe,
        semantic_hint,
        refresh_token="",
    ):
        raw_items = self._normalize_items(items_list)
        clean_items = [
            self._selected_as_string(item).strip()
            for item in raw_items
            if self._selected_as_string(item).strip()
        ]

        selection_value = str(self._unwrap_scalar(selection, "0") or "0")
        fallback_value = str(self._unwrap_scalar(fallback, "") or "")
        clamp_value = bool(self._unwrap_scalar(clamp_indexes, True))
        dedupe_value = bool(self._unwrap_scalar(dedupe, True))
        semantic_value = str(self._unwrap_scalar(semantic_hint, "document_id") or "document_id")
        refresh = _dirty_token("document_index_selector", refresh_token, selection_value, len(clean_items))

        if not clean_items:
            selected = [fallback_value] if fallback_value else []
            first = selected[0] if selected else ""
            doc_id, field_path, file_path, filename, text = self._semantic_outputs(first, semantic_value)
            status = "Input list was empty. Returned fallback." if fallback_value else "Input list was empty."
            return (
                selected,
                [doc_id] if doc_id else [],
                first,
                doc_id,
                field_path,
                file_path,
                filename,
                text,
                _json_text(selected),
                _indexed_list_text(selected),
                len(selected),
                bool(selected),
                status,
            )

        indexes = self._parse_selection_expression(
            selection=selection_value,
            item_count=len(clean_items),
            clamp_indexes=clamp_value,
        )
        if dedupe_value:
            indexes = self._dedupe_indexes(indexes)

        selected = [clean_items[idx] for idx in indexes if 0 <= idx < len(clean_items)]

        if not selected and fallback_value:
            selected = [fallback_value]

        first = selected[0] if selected else ""
        doc_id, field_path, file_path, filename, text = self._semantic_outputs(first, semantic_value)

        document_id_links = []
        for item in selected:
            item_doc_id, _, _, _, _ = self._semantic_outputs(item, "document_id")
            if item_doc_id:
                document_id_links.append(item_doc_id)

        selected_json = _json_text(selected)
        indexed_selected = _indexed_list_text(selected)
        status = (
            f"Selected {len(selected)} item(s) from {len(clean_items)} using selection={selection_value!r}; "
            f"indexes={indexes}; first={first or '<empty>'}; refresh={refresh}"
        )

        return (
            selected,
            document_id_links,
            first,
            doc_id,
            field_path,
            file_path,
            filename,
            text,
            selected_json,
            indexed_selected,
            len(selected),
            bool(selected),
            status,
        )

NODE_CLASS_MAPPINGS = {
    "ZMongoDocumentIndexSelector": ZMongoDocumentIndexSelector,
    "ZMongoDocumentSelectNthItem": ZMongoDocumentSelectNthItem,
    "ZMongoDocumentFilePathBrowser": ZMongoDocumentFilePathBrowser,
    "ZMongoDocumentHealth": ZMongoDocumentHealth,
    "ZMongoDocumentWhoAmI": ZMongoDocumentWhoAmI,
    "ZMongoDocumentUploadFile": ZMongoDocumentUploadFile,
    "ZMongoDocumentIngestTextFile": ZMongoDocumentIngestTextFile,
    "ZMongoDocumentCreateText": ZMongoDocumentCreateText,
    "ZMongoDocumentList": ZMongoDocumentList,
    "ZMongoDocumentGet": ZMongoDocumentGet,
    "ZMongoDocumentGetText": ZMongoDocumentGetText,
    "ZMongoDocumentFieldPaths": ZMongoDocumentFieldPaths,
    "ZMongoDocumentGetValue": ZMongoDocumentGetValue,
    "ZMongoDocumentSaveText": ZMongoDocumentSaveText,
    "ZMongoDocumentSaveValue": ZMongoDocumentSaveValue,
    "ZMongoDocumentUpdateMetadata": ZMongoDocumentUpdateMetadata,
    "ZMongoDocumentExtractText": ZMongoDocumentExtractText,
    "ZMongoDocumentQueueOCR": ZMongoDocumentQueueOCR,
    "ZMongoDocumentOCRStatus": ZMongoDocumentOCRStatus,
    "ZMongoDocumentDelete": ZMongoDocumentDelete,
}

NODE_DISPLAY_NAME_MAPPINGS = {
    "ZMongoDocumentIndexSelector": "06 Document Index Selector",
    "ZMongoDocumentSelectNthItem": "06 Select Nth Document Item",
    "ZMongoDocumentFilePathBrowser": "06 Document File Browser",
    "ZMongoDocumentHealth": "06 Document API Health",
    "ZMongoDocumentWhoAmI": "06 Document WhoAmI",
    "ZMongoDocumentUploadFile": "06 Upload Document File",
    "ZMongoDocumentIngestTextFile": "06 Ingest Text PDF / DOCX",
    "ZMongoDocumentCreateText": "06 Create Text Document",
    "ZMongoDocumentList": "06 List Documents",
    "ZMongoDocumentGet": "06 Get Document",
    "ZMongoDocumentGetText": "06 Get Document Text",
    "ZMongoDocumentFieldPaths": "06 Document Field Paths",
    "ZMongoDocumentGetValue": "06 Get Document Value",
    "ZMongoDocumentSaveText": "06 Save Document Text",
    "ZMongoDocumentSaveValue": "06 Save Document Value",
    "ZMongoDocumentUpdateMetadata": "06 Update Document Metadata",
    "ZMongoDocumentExtractText": "06 Extract Document Text",
    "ZMongoDocumentQueueOCR": "06 Queue Document OCR",
    "ZMongoDocumentOCRStatus": "06 Document OCR Status",
    "ZMongoDocumentDelete": "06 Delete Document",
}

__all__ = ["NODE_CLASS_MAPPINGS", "NODE_DISPLAY_NAME_MAPPINGS"]